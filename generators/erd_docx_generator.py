import os
import copy
import shutil
import subprocess
from pathlib import Path
from datetime import date
from typing import Dict, Any

from docx import Document
from docx.shared import Inches

try:
    from dotenv import load_dotenv
except ImportError:
    def load_dotenv():
        return None

load_dotenv()

TEMPLATE_PATH = os.getenv("ERD_TEMPLATE_PATH", "./template/엔티티 관계 모형 설계서.docx")
OUTPUT_PATH = os.getenv("ERD_OUTPUT_PATH", f"./output/ERD_설계서_{date.today()}.docx")


def clean_text(text):
    if text is None:
        return ""
    return str(text).strip()


def set_cell(cell, value):
    cell.text = clean_text(value)


def make_mermaid_column_type(col):
    col_type = col.get("type", "VARCHAR")
    length = col.get("length", "")
    if length:
        col_type = f"{col_type}_{length}"
    return "".join(ch if ch.isalnum() or ch == "_" else "_" for ch in col_type)


def generate_mermaid_erd(erd: Dict[str, Any]) -> str:
    lines = ["erDiagram"]

    for entity in erd.get("entities", []):
        entity_name = entity.get("entity_name", "UNKNOWN")
        lines.append(f"    {entity_name} {{")

        for col in entity.get("columns", []):
            col_type = make_mermaid_column_type(col)
            col_name = col.get("name", "UNKNOWN_COL")

            keys = []
            if col.get("pk") == "Y":
                keys.append("PK")
            if col.get("fk") == "Y":
                keys.append("FK")

            key_text = " ".join(keys)
            if key_text:
                lines.append(f"        {col_type} {col_name} {key_text}")
            else:
                lines.append(f"        {col_type} {col_name}")

        lines.append("    }")

    for rel in erd.get("relationships", []):
        from_entity = rel.get("from_entity", "")
        to_entity = rel.get("to_entity", "")
        relationship = rel.get("relationship", "")
        desc = rel.get("description", "")

        if relationship == "1:N":
            mermaid_rel = "||--o{"
        elif relationship == "1:1":
            mermaid_rel = "||--||"
        elif relationship == "N:M":
            mermaid_rel = "}o--o{"
        else:
            mermaid_rel = "||--o{"

        if from_entity and to_entity:
            lines.append(f'    {from_entity} {mermaid_rel} {to_entity} : "{desc}"')

    return "\n".join(lines)


def save_mermaid_files(erd: Dict[str, Any]):
    mmd_dir = Path("./mmd_temp")
    mmd_dir.mkdir(exist_ok=True)

    mmd_path = mmd_dir / "erd_diagram.mmd"
    png_path = mmd_dir / "erd_diagram.png"

    mermaid_code = generate_mermaid_erd(erd)
    mmd_path.write_text(mermaid_code, encoding="utf-8")

    mmdc_path = os.getenv("MMDC_PATH") or shutil.which("mmdc") or r"C:\Users\Playdata\AppData\Roaming\npm\mmdc.cmd"
    env = os.environ.copy()
    puppeteer_executable_path = env.get("PUPPETEER_EXECUTABLE_PATH") or find_puppeteer_browser()
    if puppeteer_executable_path:
        env["PUPPETEER_EXECUTABLE_PATH"] = puppeteer_executable_path

    try:
        subprocess.run(
            [
                mmdc_path,
                "-i", str(mmd_path),
                "-o", str(png_path),
                "-b", "white",
            ],
            check=True,
            env=env,
        )
    except (FileNotFoundError, subprocess.CalledProcessError) as e:
        print(f"[WARN] Mermaid 이미지 생성 실패: {e}")
        print("[WARN] mmd_temp/erd_diagram.mmd 파일만 저장하고 ERD 이미지 삽입은 건너뜁니다.")
        return str(mmd_path), None

    return str(mmd_path), str(png_path)


def find_puppeteer_browser():
    cache_dir = Path.home() / ".cache" / "puppeteer"
    candidates = []
    patterns = [
        "chrome-headless-shell/**/chrome-headless-shell.exe",
        "chrome/**/chrome.exe",
    ]

    for pattern in patterns:
        candidates.extend(path for path in cache_dir.glob(pattern) if path.is_file())

    if not candidates:
        return None

    candidates.sort(key=lambda path: path.stat().st_mtime, reverse=True)
    return str(candidates[0])


def fill_header_table(doc, erd):
    table = doc.tables[0]
    set_cell(table.cell(1, 1), erd.get("system_name", ""))
    set_cell(table.cell(2, 1), erd.get("stage_name", "설계"))
    set_cell(table.cell(2, 4), erd.get("created_date", str(date.today())))
    set_cell(table.cell(2, 6), erd.get("version", "v1.0"))


def fill_erd_table(doc, erd):
    table = doc.tables[1]
    set_cell(table.cell(0, 1), erd.get("erd_id", ""))
    set_cell(table.cell(0, 3), erd.get("erd_name", ""))

    rel_text = []
    for rel in erd.get("relationships", []):
        rel_text.append(
            f'{rel.get("from_entity", "")} {rel.get("relationship", "")} '
            f'{rel.get("to_entity", "")} - {rel.get("description", "")}'
        )

    if len(table.rows) > 1:
        set_cell(table.cell(1, 0), "\n".join(rel_text))


def insert_erd_image(doc, image_path, erd):
    if len(doc.tables) > 1 and len(doc.tables[1].rows) > 1:
        cell = doc.tables[1].cell(1, 0)
        rel_text = cell.text
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(6.5))

        if rel_text.strip():
            paragraph = cell.add_paragraph()
            paragraph.add_run(rel_text)
        return

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(6.5))


def fill_entity_tables_fast(doc, erd):
    entities = erd.get("entities", [])
    if not entities:
        return

    table = doc.tables[2]

    set_cell(table.cell(0, 2), "ALL")
    set_cell(table.cell(0, 7), "전체 엔티티 명세")
    set_cell(table.cell(1, 4), "요구사항 기반으로 도출된 전체 엔티티 및 속성 목록")

    base_row_idx = 3
    first = True

    for entity in entities:
        if first:
            row = table.rows[base_row_idx]
            first = False
        else:
            row = table.add_row()

        values = [
            f"[{entity.get('entity_name', '')}]",
            entity.get("entity_description", ""),
            "", "", "", "", "", "", "", "",
        ]
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)

        for col in entity.get("columns", []):
            row = table.add_row()
            values = [
                col.get("name", ""),
                col.get("synonym", ""),
                col.get("type", ""),
                col.get("length", ""),
                col.get("not_null", ""),
                col.get("pk", ""),
                col.get("fk", ""),
                col.get("inx", ""),
                col.get("default", ""),
                col.get("constraint", ""),
            ]
            for cell, value in zip(row.cells, values):
                set_cell(cell, value)


def clone_table_after(table):
    new_tbl = copy.deepcopy(table._tbl)
    table._tbl.addnext(new_tbl)


def fill_entity_table(table, entity):
    set_cell(table.cell(0, 2), entity.get("entity_id", ""))
    set_cell(table.cell(0, 7), entity.get("entity_name", ""))
    set_cell(table.cell(1, 4), entity.get("entity_description", ""))

    base_row_idx = 3
    for idx, col in enumerate(entity.get("columns", [])):
        if idx == 0:
            row = table.rows[base_row_idx]
        else:
            row = table.add_row()

        values = [
            col.get("name", ""),
            col.get("synonym", ""),
            col.get("type", ""),
            col.get("length", ""),
            col.get("not_null", ""),
            col.get("pk", ""),
            col.get("fk", ""),
            col.get("inx", ""),
            col.get("default", ""),
            col.get("constraint", ""),
        ]
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)


def fill_entity_tables_separate(doc, erd):
    entities = erd.get("entities", [])
    if not entities:
        return

    template_table = doc.tables[2]
    for _ in entities[1:]:
        clone_table_after(template_table)

    entity_tables = doc.tables[2:2 + len(entities)]
    for table, entity in zip(entity_tables, entities):
        fill_entity_table(table, entity)


def save_docx_with_fallback(doc, output_path: str) -> str:
    path = Path(output_path)

    for idx in range(0, 100):
        candidate = path if idx == 0 else path.with_name(f"{path.stem}_{idx}{path.suffix}")
        try:
            doc.save(str(candidate))
            return str(candidate)
        except PermissionError:
            continue

    raise PermissionError(f"DOCX 저장 실패: {output_path} 및 대체 파일명을 사용할 수 없습니다.")


def generate_erd_docx(
    erd: Dict[str, Any],
    template_path: str = TEMPLATE_PATH,
    output_path: str = OUTPUT_PATH,
    *,
    use_mermaid: bool = True,
    fast_table: bool = False,
):
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc = Document(template_path)

    fill_header_table(doc, erd)
    fill_erd_table(doc, erd)

    if use_mermaid:
        print("[DOCX] Mermaid ERD 이미지 생성")
        mmd_path, png_path = save_mermaid_files(erd)
        if png_path:
            print("[DOCX] ERD 이미지 삽입")
            insert_erd_image(doc, png_path, erd)
        print("[완료] Mermaid code:", mmd_path)
        if png_path:
            print("[완료] Mermaid image:", png_path)

    if fast_table:
        fill_entity_tables_fast(doc, erd)
    else:
        fill_entity_tables_separate(doc, erd)

    saved_path = save_docx_with_fallback(doc, output_path)
    if saved_path != output_path:
        print(f"[WARN] 기존 파일을 덮어쓸 수 없어 다른 이름으로 저장했습니다: {saved_path}")
    print(f"[완료] ERD 설계서: {saved_path}")
    return saved_path


if __name__ == "__main__":
    import json

    with open("./output/erd_agent_output.json", "r", encoding="utf-8") as f:
        erd = json.load(f)

    generate_erd_docx(erd)
