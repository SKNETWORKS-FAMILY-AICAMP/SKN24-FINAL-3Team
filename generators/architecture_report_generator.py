from pathlib import Path

from docx import Document
from docx.shared import Inches

from agents.arch_nodes.common import strip_mermaid_block
from generators.architecture_image_generator import render_mermaid_image


def _add_markdownish_content(doc: Document, content: str):
    for raw_line in str(content or "").splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)


def generate_architecture_docx(
    *,
    report_specs: str,
    mermaid_script: str,
    image_path: str | None,
    output_docx_path: str,
) -> str:
    output_path = Path(output_docx_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("아키텍처 설계서", level=1)
    _add_markdownish_content(doc, report_specs)

    doc.add_heading("시스템 아키텍처 다이어그램", level=2)
    if image_path and Path(image_path).exists():
        doc.add_picture(image_path, width=Inches(6.5))

    doc.add_heading("Mermaid 원본", level=2)
    doc.add_paragraph(strip_mermaid_block(mermaid_script))

    doc.save(str(output_path))
    return str(output_path)


def generate_architecture_report(
    *,
    report_specs: str,
    mermaid_script: str,
    output_md_path: str,
    output_docx_path: str,
    output_image_path: str,
    render_image: bool = True,
) -> dict:
    md_path = Path(output_md_path)
    md_path.parent.mkdir(parents=True, exist_ok=True)

    image_path = None
    mmd_path = None
    if render_image:
        mmd_path, image_path = render_mermaid_image(
            mermaid_script,
            output_image_path=output_image_path,
        )

    image_section = ""
    if image_path:
        image_section = f"\n\n## 시스템 아키텍처 다이어그램\n\n![Architecture]({Path(image_path).as_posix()})"

    content = (
        f"{report_specs}\n"
        f"{image_section}\n\n"
        f"## Mermaid 원본\n\n"
        f"```mermaid\n{strip_mermaid_block(mermaid_script)}\n```\n"
    )
    md_path.write_text(content, encoding="utf-8")

    docx_path = generate_architecture_docx(
        report_specs=report_specs,
        mermaid_script=mermaid_script,
        image_path=image_path,
        output_docx_path=output_docx_path,
    )

    return {
        "md_path": str(md_path),
        "docx_path": docx_path,
        "mmd_path": mmd_path,
        "image_path": image_path,
    }

