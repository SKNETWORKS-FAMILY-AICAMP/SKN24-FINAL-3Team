from pathlib import Path
from typing import Any


def generate_ts_docx(data: dict[str, Any], output_path: str) -> str:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    def set_cell_bg(cell, hex_color: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    def add_text(cell, text: Any, *, center: bool = False, is_auto: bool = False) -> None:
        paragraph = cell.paragraphs[0]
        if center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("" if text is None else str(text))
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(8)
        if is_auto:
            run.font.color.rgb = RGBColor(0xD9, 0x00, 0x00)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def add_header(table, headers: list[str]) -> None:
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            set_cell_bg(cell, "4472C4")
            add_text(cell, header, center=True)
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)

    title = doc.add_heading("통합시험 시나리오", level=1)
    title.runs[0].font.name = "Malgun Gothic"
    title.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    cases_by_scenario: dict[str, list[dict[str, Any]]] = {}
    for case in data.get("cases", []):
        cases_by_scenario.setdefault(case.get("scenario_id", ""), []).append(case)

    for scenario in data.get("scenarios", []):
        doc.add_heading(f"{scenario.get('scenario_id', '')} - {scenario.get('scenario_name', '')}", level=2)
        doc.add_paragraph(scenario.get("scenario_description", ""))

        scenario_headers = [
            "시험시나리오 ID",
            "시험시나리오명",
            "시험케이스 ID",
            "시험케이스 설명",
            "시험 절차",
            "시나리오 설명",
            "비고",
        ]
        scenario_table = doc.add_table(rows=1, cols=len(scenario_headers))
        scenario_table.style = "Table Grid"
        add_header(scenario_table, scenario_headers)

        for test_case in scenario.get("test_cases", []):
            row = scenario_table.add_row().cells
            add_text(row[0], scenario.get("scenario_id", ""), center=True)
            add_text(row[1], scenario.get("scenario_name", ""))
            add_text(row[2], test_case.get("test_case_id", ""), center=True)
            add_text(row[3], test_case.get("test_case_description", ""))
            add_text(
                row[4],
                "\n".join(f"{idx + 1}. {step}" for idx, step in enumerate(test_case.get("test_procedure", []))),
            )
            add_text(row[5], test_case.get("scenario_detail", ""))
            add_text(row[6], test_case.get("note", ""))

        case_headers = [
            "차수",
            "시험시나리오 ID",
            "시험케이스 ID",
            "순번",
            "처리 내용",
            "시험 항목",
            "사전조건",
            "입력값",
            "예상 결과",
            "화면 ID",
            "시험결과",
            "비고",
        ]
        case_table = doc.add_table(rows=1, cols=len(case_headers))
        case_table.style = "Table Grid"
        add_header(case_table, case_headers)

        for case in cases_by_scenario.get(scenario.get("scenario_id", ""), []):
            row = case_table.add_row().cells
            is_auto = "자동 보완" in str(case.get("note", ""))
            values = [
                case.get("round"),
                case.get("scenario_id"),
                case.get("test_case_id"),
                case.get("sequence"),
                case.get("process_content"),
                case.get("test_item"),
                case.get("precondition"),
                case.get("input_data"),
                case.get("expected_result"),
                case.get("screen_id"),
                case.get("test_result"),
                case.get("note"),
            ]
            for index, value in enumerate(values):
                add_text(row[index], value, center=index in {0, 1, 2, 3, 9}, is_auto=is_auto)

        doc.add_page_break()

    doc.save(output)
    return str(output)
