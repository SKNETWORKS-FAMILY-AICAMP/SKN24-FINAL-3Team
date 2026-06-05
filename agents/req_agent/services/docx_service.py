"""
services/docx_service.py
두 개의 분리된 표(메타 표, 데이터 표) 구조를 정확히 인식하여 
템플릿에 데이터를 매핑하는 서비스 (런팟 한글 깨짐 방지 포함)
"""
import os
import logging
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# RunPod 최상위 workspace 절대 경로 지정
TEMPLATE_PATH = "/workspace/template/사용자 요구사항 명세서.docx"
_OUT_DIR = Path(__file__).parent.parent / "output"

_STATUS_COLOR = {
    "신규": RGBColor(0xE8, 0xF5, 0xE9),
    "수정": RGBColor(0xE3, 0xF2, 0xFD),
    "기존": RGBColor(0xFF, 0xFF, 0xFF),
}

def generate_docx(reqs: list[dict], prefix: str = "요구사항",
                  output_path: str = None) -> str:
    _OUT_DIR.mkdir(exist_ok=True)
    ts  = datetime.now().strftime("%Y%m%d_%H%M%S")
    out = Path(output_path) if output_path else _OUT_DIR / f"{prefix}_{ts}.docx"

    clean = [{k: v for k, v in r.items() if not k.startswith("_")} for r in reqs]
    
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"지정하신 템플릿 파일을 찾을 수 없습니다: {TEMPLATE_PATH}")
    
    doc = Document(TEMPLATE_PATH)
    doc.settings.language = 'ko-KR'

    # [교정 핵심] 문서 내에 표가 정상적으로 2개 이상 존재하는지 검증
    if len(doc.tables) < 2:
        raise ValueError("템플릿에 표가 부족합니다. 메타 표와 데이터 표 2개가 필요합니다.")

    meta_table = doc.tables[0]  # 첫 번째 표: 시스템명, 단계명, 작성일자 등
    data_table = doc.tables[1]  # 두 번째 표: 요구사항 ID, 요구사항 명 등 (진짜 데이터판)

    # 1. 상단 메타 표에 날짜 및 버전 매핑
    _fill_meta_table(meta_table)

    # 2. 하단 데이터 표에 요구사항 리스트 정밀 매핑
    _fill_data_table(data_table, clean)

    doc.save(str(out))
    _fix_xml(out)
    logger.info("docx 템플릿 정밀 매핑 및 저장 완료: %s", out)
    return str(out)


def _fill_meta_table(table):
    """첫 번째 표(메타 정보)의 정확한 좌표에 값을 삽입"""
    try:
        today_str = datetime.now().strftime("%Y-%m-%d")
        # 행2, 열3 -> '작성일자' 우측 빈칸에 오늘 날짜 입력
        _set_cell_with_font(table.cell(2, 3), today_str, bold=False, center=True)
        # 행2, 열5 -> '버전' 우측 빈칸에 버전 입력
        _set_cell_with_font(table.cell(2, 5), "v1.0", bold=False, center=True)
    except Exception as e:
        logger.warning(f"메타 표 채우기 중 일부 인덱스 건너뜀: {e}")


def _fill_data_table(table, reqs):
    """두 번째 표(요구사항 목록)의 헤더 아래 영역에 데이터를 채움"""
    # 표 2의 0번째 행은 [요구사항 ID, 요구사항 명...] 헤더이므로 데이터는 1번째 행부터 시작
    base_row_idx = 1 
    
    for i, req in enumerate(reqs):
        # 템플릿에 이미 존재하는 빈 행(기본 1줄)이 있으면 재사용하고, 부족하면 새로 생성
        if base_row_idx + i < len(table.rows):
            row = table.rows[base_row_idx + i]
        else:
            row = table.add_row()
            
        bg = _STATUS_COLOR.get(req.get("status","기존"), _STATUS_COLOR["기존"])
        src  = _join(req.get("source"))
        cons = _join(req.get("constraints"))
        crit = _join(req.get("validation_criteria"))
        
        # 템플릿 컬럼 서식 순서대로 (총 10개 컬럼) 정확히 매핑
        vals = [
            req.get("requirement_id",""),    req.get("requirement_name",""),
            req.get("requirement_type",""),  req.get("description",""),
            src, cons, req.get("priority",""),
            req.get("note","") or "",        crit,
            req.get("status",""),
        ]
        
        centers = {0, 2, 6, 9}  # ID, 구분, 중요도, 상태는 가운데 정렬
        for j, v in enumerate(vals):
            if j < len(row.cells):
                _set_cell_with_font(row.cells[j], str(v), bg=bg, center=(j in centers))


def _set_cell_with_font(cell, text, bold=False, bg=None, center=False, size=10, font_name="맑은 고딕"):
    """글자를 채우면서 리눅스 환경으로 인해 폰트가 깨지는 현상(네모박스)을 XML 레벨에서 원천 차단"""
    cell.text = ""  
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    
    # XML에 동아시아 한글 폰트 구조 강제 주입
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)  
    rPr.append(rFonts)
    
    lang = OxmlElement('w:lang')
    lang.set(qn('w:eastAsia'), 'ko-KR')
    rPr.append(lang)
    
    if bg:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), f"{bg[0]:02X}{bg[1]:02X}{bg[2]:02X}")
        tcPr.insert(0, shd)


def _join(val):
    if isinstance(val, list): return "\n".join(str(v) for v in val if v)
    return str(val or "")


def _fix_xml(path: Path):
    import zipfile, re, shutil
    tmp = path.parent / (path.stem + "_tmp.docx")
    with zipfile.ZipFile(path, 'r') as zin,           zipfile.ZipFile(tmp,  'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/settings.xml':
                text = data.decode('utf-8')
                text = re.sub(r'<w:zoom([^>]*/?)>', lambda m: f'<w:zoom{m.group(1)}>' if 'percent' in m.group(1) else '<w:zoom w:percent="100" w:val="bestFit"/>', text)
                data = text.encode('utf-8')
            elif item.filename == 'word/document.xml':
                text = data.decode('utf-8')
                def fix_tcpr(m):
                    tcpr = m.group(0)
                    tcw_m = re.search(r'<w:tcW[^/]*/>', tcpr)
                    if not tcw_m: return tcpr
                    tcw = tcw_m.group(0)
                    tcpr = tcpr.replace(tcw, '')
                    return tcpr.replace('<w:tcPr>', f'<w:tcPr>{tcw}', 1)
                text = re.sub(r'<w:tcPr>.*?</w:tcPr>', fix_tcpr, text, flags=re.S)
                data = text.encode('utf-8')
            zout.writestr(item, data)
    shutil.move(str(tmp), str(path))