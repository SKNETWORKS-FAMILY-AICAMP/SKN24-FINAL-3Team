import json
import os
import re
from datetime import date
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document

from rag.rag_service import search_rag

try:
    from dotenv import load_dotenv
except ImportError:
    def load_dotenv():
        return None

load_dotenv()

ERD_DOCX_PATH = os.getenv("ERD_DOCX_PATH", f"./output/ERD_설계서_{date.today()}.docx")
OUTPUT_JSON_PATH = os.getenv("DB_DESIGN_JSON_PATH", "./json_temp/database_design_agent_output.json")


def clean_text(value: Any) -> str:
    return "" if value is None else str(value).strip()


def cell_text(row, idx: int) -> str:
    if idx >= len(row.cells):
        return ""
    return clean_text(row.cells[idx].text)


def find_latest_erd_docx(output_dir: str = "./output") -> Optional[str]:
    paths = sorted(
        Path(output_dir).glob("ERD_설계서*.docx"),
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )
    return str(paths[0]) if paths else None


def resolve_erd_docx_path(erd_docx_path: Optional[str] = None) -> str:
    if erd_docx_path:
        if Path(erd_docx_path).exists():
            return erd_docx_path
        raise FileNotFoundError(f"지정한 ERD 설계서 DOCX를 찾지 못했습니다: {erd_docx_path}")

    if Path(ERD_DOCX_PATH).exists():
        return ERD_DOCX_PATH

    latest_path = find_latest_erd_docx()
    if latest_path:
        return latest_path

    requested = erd_docx_path or ERD_DOCX_PATH
    raise FileNotFoundError(f"ERD 설계서 DOCX를 찾지 못했습니다: {requested}")


def make_database_id(system_name: str, erd_id: str) -> str:
    if erd_id:
        return "DB-" + re.sub(r"^ERD-", "", erd_id)
    if system_name:
        token = re.sub(r"[^A-Za-z0-9]+", "_", system_name).strip("_").upper()
        return f"DB-{token[:30]}"
    return "DB-001"


def normalize_type(column: Dict[str, str]) -> str:
    col_type = column.get("type", "")
    length = column.get("length", "")
    if length:
        return f"{col_type}({length})"
    return col_type


def parse_relationships(text: str) -> List[Dict[str, str]]:
    relationships = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue

        match = re.match(r"(.+?)\s+(1:1|1:N|N:M)\s+(.+?)\s+-\s+(.+)", line)
        if match:
            relationships.append({
                "from_entity": match.group(1).strip(),
                "relationship": match.group(2).strip(),
                "to_entity": match.group(3).strip(),
                "description": match.group(4).strip(),
            })
    return relationships


def parse_column_row(row) -> Dict[str, str]:
    return {
        "name": cell_text(row, 0),
        "synonym": cell_text(row, 1),
        "type": cell_text(row, 2),
        "length": cell_text(row, 3),
        "not_null": cell_text(row, 4),
        "pk": cell_text(row, 5),
        "fk": cell_text(row, 6),
        "inx": cell_text(row, 7),
        "default": cell_text(row, 8),
        "constraint": cell_text(row, 9),
    }


def parse_separate_entity_table(table) -> Dict[str, Any]:
    entity = {
        "entity_id": cell_text(table.rows[0], 2),
        "entity_name": cell_text(table.rows[0], 7),
        "entity_description": cell_text(table.rows[1], 4),
        "columns": [],
    }

    header_idx = 2
    for idx, row in enumerate(table.rows):
        if cell_text(row, 0) in ("속성명", "컬럼명"):
            header_idx = idx
            break

    for row in table.rows[header_idx + 1:]:
        if not cell_text(row, 0):
            continue
        entity["columns"].append(parse_column_row(row))

    return entity


def parse_fast_entity_table(table) -> List[Dict[str, Any]]:
    entities = []
    current = None

    for row in table.rows[3:]:
        first = cell_text(row, 0)
        if not first:
            continue

        entity_match = re.match(r"^\[(.+)\]$", first)
        if entity_match:
            current = {
                "entity_id": f"ENT-{len(entities) + 1:03d}",
                "entity_name": entity_match.group(1),
                "entity_description": cell_text(row, 1),
                "columns": [],
            }
            entities.append(current)
            continue

        if current:
            current["columns"].append(parse_column_row(row))

    return entities


def parse_erd_docx(erd_docx_path: str) -> Dict[str, Any]:
    doc = Document(erd_docx_path)
    if len(doc.tables) < 3:
        raise ValueError("ERD 설계서 DOCX에서 필요한 표를 찾지 못했습니다.")

    header_table = doc.tables[0]
    erd_table = doc.tables[1]

    erd = {
        "system_name": cell_text(header_table.rows[1], 1),
        "stage_name": cell_text(header_table.rows[2], 1) or "설계",
        "created_date": cell_text(header_table.rows[2], 4) or str(date.today()),
        "version": cell_text(header_table.rows[2], 6) or "v1.0",
        "erd_id": cell_text(erd_table.rows[0], 1),
        "erd_name": cell_text(erd_table.rows[0], 3),
        "relationships": parse_relationships(cell_text(erd_table.rows[1], 0) if len(erd_table.rows) > 1 else ""),
        "entities": [],
    }

    for table in doc.tables[2:]:
        if len(table.rows) < 3 or cell_text(table.rows[0], 0) != "엔티티 ID":
            continue

        if cell_text(table.rows[0], 2) == "ALL":
            erd["entities"].extend(parse_fast_entity_table(table))
        else:
            erd["entities"].append(parse_separate_entity_table(table))

    return erd


def build_database_design(erd: Dict[str, Any]) -> Dict[str, Any]:
    system_name = erd.get("system_name", "")
    database_id = make_database_id(system_name, erd.get("erd_id", ""))
    database_name = f"{system_name} 데이터베이스" if system_name else "업무 데이터베이스"

    tables = []
    for idx, entity in enumerate(erd.get("entities", []), start=1):
        table_id = entity.get("entity_id") or f"TBL-{idx:03d}"
        table_name = entity.get("entity_name", "")
        tables.append({
            "table_id": table_id,
            "table_name": table_name,
            "database_name": database_name,
            "tablespace_name": f"TS_{table_name}"[:30] if table_name else f"TS_{idx:03d}",
            "trigger_config": "해당 없음",
            "table_description": entity.get("entity_description", ""),
            "initial_count": "0",
            "daily_growth": "산정 필요",
            "retention_period": "업무 기준에 따름",
            "max_count": "산정 필요",
            "capacity": "산정 필요",
            "note": "",
            "columns": [
                {
                    "column_name": col.get("synonym", "") or col.get("name", ""),
                    "column_id": col.get("name", ""),
                    "type_and_length": normalize_type(col),
                    "not_null": col.get("not_null", ""),
                    "pk": col.get("pk", ""),
                    "fk": col.get("fk", ""),
                    "idx": col.get("inx", ""),
                    "default": col.get("default", ""),
                    "constraint": col.get("constraint", ""),
                }
                for col in entity.get("columns", [])
            ],
        })

    return {
        "system_name": system_name,
        "subsystem_name": erd.get("erd_name", ""),
        "stage_name": erd.get("stage_name", "설계"),
        "created_date": erd.get("created_date", str(date.today())),
        "version": erd.get("version", "v1.0"),
        "databases": [{
            "database_id": database_id,
            "database_name": database_name,
            "owner_department": "업무 담당 부서",
            "note": erd.get("erd_name", ""),
            "storage_group": "업무 기준에 따름",
            "bufferpool": "업무 기준에 따름",
            "index_bufferpool": "업무 기준에 따름",
        }],
        "tables": tables,
        "relationships": erd.get("relationships", []),
    }


def append_constraint(constraint: str, addition: str) -> str:
    constraint = clean_text(constraint)
    addition = clean_text(addition)

    if not addition or addition in constraint:
        return constraint
    if not constraint:
        return addition
    return f"{constraint}; {addition}"


def has_any_keyword(text: str, keywords: List[str]) -> bool:
    upper_text = text.upper()
    return any(keyword.upper() in upper_text for keyword in keywords)


def infer_privacy_and_encryption_note(column: Dict[str, str]) -> str:
    text = " ".join([
        column.get("column_name", ""),
        column.get("column_id", ""),
        column.get("constraint", ""),
    ])

    notes = []
    if has_any_keyword(text, ["PASSWORD", "PWD", "PASS", "비밀번호"]):
        notes.append("개인정보/인증정보 여부 검토")
        notes.append("암호화 또는 해시 저장 대상")
    elif has_any_keyword(text, ["ACCOUNT_NO", "ACCOUNT_NUMBER", "계좌번호", "PHONE", "TEL", "EMAIL", "RRN", "주민등록", "고객명", "성명"]):
        notes.append("개인정보 포함 여부 검토")
        notes.append("암호화 여부 검토")

    return "; ".join(notes)


def build_database_design_rag_context(design: Dict[str, Any]) -> Dict[str, Any]:
    queries = [
        "컬럼정의서 Not Null 여부 PK 정보 FK 정보 제약조건 개인정보 여부 암호화 여부 작성지침",
        "테이블정의서 보존기간 테이블볼륨 발생주기 작성지침",
        "데이터베이스정의서 데이터용량 작성지침",
    ]

    for table in design.get("tables", []):
        column_ids = " ".join(column.get("column_id", "") for column in table.get("columns", []))
        queries.append(
            f"{table.get('table_name', '')} {column_ids} PK FK Not Null 제약조건 개인정보 암호화"
        )

    rows = []
    seen_chunk_ids = set()
    for query in queries:
        for row in search_rag(
            query=query,
            domain="public_data",
            applies_to="database_design",
            doc_type="db_standard_manual",
            limit=4,
        ):
            metadata = row.get("metadata", {})
            chunk_id = metadata.get("chunk_id") or metadata.get("title") or row.get("text", "")[:80]
            if chunk_id in seen_chunk_ids:
                continue
            seen_chunk_ids.add(chunk_id)
            rows.append(row)

    return {
        "source": "공공데이터베이스 표준화 관리 매뉴얼",
        "purpose": "PK/FK/Not Null/제약조건/개인정보/암호화 여부 보강",
        "rows": rows,
    }


def compact_database_design_rag_context(rag_context: Dict[str, Any]) -> Dict[str, Any]:
    compact_rows = []
    for row in rag_context.get("rows", []):
        metadata = row.get("metadata", {})
        compact_rows.append({
            "score": row.get("score"),
            "title": metadata.get("title"),
            "section": metadata.get("section"),
            "page": metadata.get("page"),
            "chunk_type": metadata.get("chunk_type"),
            "text": row.get("text", "")[:900],
        })

    return {
        "source": rag_context.get("source", ""),
        "purpose": rag_context.get("purpose", ""),
        "rows": compact_rows,
    }


def enhance_database_design_with_rag(design: Dict[str, Any], rag_context: Dict[str, Any]) -> Dict[str, Any]:
    compact_rows = compact_database_design_rag_context(rag_context).get("rows", [])
    evidence_pages = sorted({
        str(row.get("page"))
        for row in compact_rows
        if row.get("page")
    })
    evidence_note = f"공공DB 표준화 관리 매뉴얼 RAG 근거 p.{', '.join(evidence_pages[:6])}" if evidence_pages else "공공DB 표준화 관리 매뉴얼 RAG 근거"

    for table in design.get("tables", []):
        table_name = table.get("table_name", "")
        for column in table.get("columns", []):
            constraint = column.get("constraint", "")

            if column.get("pk"):
                column["not_null"] = column.get("not_null") or "Y"
                constraint = append_constraint(constraint, "PK 참여 컬럼")

            if column.get("fk"):
                constraint = append_constraint(constraint, "FK 제약 참여 컬럼")
                if table_name:
                    constraint = append_constraint(constraint, f"{table_name} 관계 무결성 검토")

            if column.get("not_null") == "Y":
                constraint = append_constraint(constraint, "입력 시 필수값")

            privacy_note = infer_privacy_and_encryption_note(column)
            if privacy_note:
                constraint = append_constraint(constraint, privacy_note)

            if constraint:
                constraint = append_constraint(constraint, evidence_note)

            column["constraint"] = constraint

    design["rag_enhancement"] = {
        "enabled": True,
        "source": rag_context.get("source", ""),
        "purpose": rag_context.get("purpose", ""),
        "evidence_pages": evidence_pages,
    }
    return design


def generate_database_design_json(
    erd_docx_path: Optional[str] = None,
    output_json_path: str = OUTPUT_JSON_PATH,
    *,
    use_rag: bool = True,
) -> Dict[str, Any]:
    resolved_path = resolve_erd_docx_path(erd_docx_path)
    print(f"[DB 설계 Agent] ERD 설계서 DOCX 읽기: {resolved_path}")
    erd = parse_erd_docx(resolved_path)
    print("[DB 설계 Agent] 데이터베이스 설계 JSON 생성")
    database_design = build_database_design(erd)

    if use_rag:
        try:
            print("[DB 설계 Agent] RAG 기반 제약조건 보강 시작")
            rag_context = build_database_design_rag_context(database_design)
            database_design = enhance_database_design_with_rag(database_design, rag_context)
            print("[DB 설계 Agent] RAG 기반 제약조건 보강 완료")
        except Exception as e:
            print(f"[WARN] DB 설계 RAG 보강 실패: {e}")

    Path(output_json_path).parent.mkdir(parents=True, exist_ok=True)
    with open(output_json_path, "w", encoding="utf-8") as f:
        json.dump(database_design, f, ensure_ascii=False, indent=2)

    print(f"[완료] DB 설계 JSON: {output_json_path}")
    return database_design


if __name__ == "__main__":
    generate_database_design_json()
