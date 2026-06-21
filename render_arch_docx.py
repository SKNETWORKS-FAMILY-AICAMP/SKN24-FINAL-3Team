"""
render_arch_docx.py
──────────────────────────
ARCH 산출물(document.json/structure.json)을 공공 ARCH 양식(arch_template.docx)에 맞춰 DOCX로 렌더링합니다.

핵심 차이점
- Document()로 새 보고서를 만들지 않고, arch_template.docx를 열어 양식에 값을 채웁니다.
- 1. 시스템 아키텍처 영역에는 Mermaid 소스와 아키텍처 그림을 생성해 삽입합니다.
  * mmdc(mermaid-cli)가 있으면 Mermaid PNG를 사용합니다.
  * mmdc가 없으면 Graphviz PNG로 자동 fallback 합니다.
- 2. 아키텍처 요구사항 및 구현방안 영역은 요구사항 내용/구현방안 표를 반복 생성합니다.

사용법 예시
    python render_arch_docx.py document.json --structure structure.json --template arch_template.docx --out architecture_fixed.docx
    python render_arch_docx.py _arch_lab/document.json --structure _arch_lab/structure.json --template templates/arch_template.docx
"""
from __future__ import annotations

import argparse
import copy
import json
import re
import shutil
import subprocess
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table

from agents.architecture_analysis.processors.diagram_builder import (
    build_clean_architecture_mermaid_source,
    select_diagram_relations,
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("document", nargs="?", default="_arch_lab/document.json", help="architecture_document_json 파일")
    parser.add_argument("--structure", default=None, help="architecture_structure_json 파일(선택)")
    parser.add_argument("--template", default="arch_template.docx", help="ARCH DOCX 템플릿")
    parser.add_argument("--out", default="_arch_lab/architecture.docx", help="출력 DOCX 경로")
    parser.add_argument("--diagram-out", default=None, help="생성할 아키텍처 다이어그램 PNG 경로")
    parser.add_argument("--mermaid-out", default=None, help="생성할 Mermaid .mmd 경로")
    parser.add_argument("--system-name", default="", help="표지 시스템명")
    parser.add_argument("--subsystem-name", default="", help="표지 서브시스템명")
    parser.add_argument("--write-date", default=date.today().isoformat(), help="작성일자")
    parser.add_argument("--version", default="v1.0", help="버전")
    return parser.parse_args()


def load_json(path: str | Path) -> dict[str, Any]:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"입력 JSON 파일이 없습니다: {p}")
    return json.loads(p.read_text(encoding="utf-8"))


def ensure_parent(path: str | Path) -> Path:
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    return p


def first_non_empty(*values: Any) -> str:
    for value in values:
        if value is not None and str(value).strip():
            return str(value).strip()
    return ""


def clean_requirement_text(text: Any) -> str:
    """RAG 중복으로 같은 근거가 2~3번 반복되는 문장을 양식 출력용으로 정리합니다."""
    s = str(text or "").strip()
    s = re.sub(r"\s+", " ", s)

    marker = "확보를 위해 "
    end_marker = "을 설계 기준"
    if marker in s and end_marker in s:
        prefix, tail = s.split(marker, 1)
        body, suffix = tail.split(end_marker, 1)
        parts = [p.strip() for p in body.split(",") if p.strip()]
        deduped: list[str] = []
        for p in parts:
            if p not in deduped:
                deduped.append(p)
        s = f"{prefix}{marker}{', '.join(deduped)}{end_marker}{suffix}"

    # 같은 [ID] 근거가 나열되는 구현방안 말미도 정리
    m = re.search(r"RAG로 확인한 비기능 근거인 (.+?)을 상세 설계 검토 기준으로 사용합니다\.", s)
    if m:
        parts = [p.strip() for p in m.group(1).split(",") if p.strip()]
        deduped: list[str] = []
        for p in parts:
            if p not in deduped:
                deduped.append(p)
        s = s[: m.start(1)] + ", ".join(deduped) + s[m.end(1) :]
    return s


def safe_node_id(value: Any, fallback: str = "NODE") -> str:
    text = str(value or fallback)
    node = re.sub(r"[^0-9A-Za-z_]+", "_", text.upper()).strip("_")
    if not node:
        node = fallback
    if node[0].isdigit():
        node = "N_" + node
    return node


def component_label(component: dict[str, Any]) -> str:
    return first_non_empty(component.get("name"), component.get("component_name"), component.get("component_id"))


def build_mermaid(data: dict[str, Any]) -> str:
    """아키텍처 다이어그램 Mermaid 소스를 생성합니다.

    단독 테스트 렌더러도 운영 agent와 같은 no-cross diagram_builder를 사용합니다.
    """
    return build_clean_architecture_mermaid_source(
        data,
        direction="LR",
        edge_label_mode="none",
        max_edges=16,
    )


def escape_mermaid(text: str) -> str:
    return str(text).replace('"', "'").replace("|", "/").replace("\n", " ")


def shorten(text: str, max_len: int) -> str:
    s = re.sub(r"\s+", " ", str(text or "")).strip()
    return s if len(s) <= max_len else s[: max_len - 1].rstrip() + "…"


def render_diagram(data: dict[str, Any], mermaid_path: Path, png_path: Path) -> Path:
    mermaid_path.write_text(build_mermaid(data), encoding="utf-8")
    ensure_parent(png_path)

    mmdc = shutil.which("mmdc")
    if mmdc:
        try:
            subprocess.run(
                [mmdc, "-i", str(mermaid_path), "-o", str(png_path), "-b", "transparent", "-w", "2000", "-H", "1150", "-s", "1.15"],
                check=True,
                timeout=60,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            if png_path.exists() and png_path.stat().st_size > 0:
                _trim_png_whitespace(png_path)
                return png_path
        except Exception:
            pass

    # Mermaid CLI가 없는 개발/CI 환경 fallback. DOCX에는 그림만 필요하므로 Graphviz로 대체 렌더링합니다.
    rendered_path = render_graphviz_diagram(data, png_path)
    _trim_png_whitespace(rendered_path)
    return rendered_path


def _trim_png_whitespace(png_path: Path, padding: int = 24) -> None:
    """Mermaid/Graphviz가 남긴 과도한 여백을 잘라 DOCX 삽입 시 글씨가 작아지는 문제를 줄입니다."""
    try:
        from PIL import Image, ImageChops
    except Exception:
        return

    try:
        img = Image.open(png_path).convert("RGBA")
        alpha = img.getchannel("A")
        bbox = alpha.getbbox()
        if bbox is None:
            bg = Image.new(img.mode, img.size, img.getpixel((0, 0)))
            diff = ImageChops.difference(img, bg)
            bbox = diff.getbbox()
        if bbox is None:
            return
        left = max(bbox[0] - padding, 0)
        upper = max(bbox[1] - padding, 0)
        right = min(bbox[2] + padding, img.width)
        lower = min(bbox[3] + padding, img.height)
        if (right - left) < img.width or (lower - upper) < img.height:
            img.crop((left, upper, right, lower)).save(png_path)
    except Exception:
        return


def render_graphviz_diagram(data: dict[str, Any], png_path: Path) -> Path:
    try:
        from graphviz import Digraph
    except Exception as exc:
        raise RuntimeError("mermaid-cli(mmdc)와 graphviz Python 패키지 중 하나가 필요합니다.") from exc

    components = [c for c in data.get("components", []) if isinstance(c, dict)]
    relations = [r for r in data.get("relations") or data.get("edges") or [] if isinstance(r, dict)]
    layers = [l for l in data.get("layers") or data.get("subgraphs") or [] if isinstance(l, dict)]
    raw_to_safe = {str(c.get("component_id")): safe_node_id(c.get("component_id")) for c in components}
    by_safe = {safe_node_id(c.get("component_id")): c for c in components}

    dot = Digraph("architecture", format="png")
    dot.attr(rankdir="LR", splines="ortho", concentrate="true", compound="true", nodesep="0.45", ranksep="0.75", margin="0.05")
    dot.attr("graph", fontname="NanumGothic")
    dot.attr("node", shape="box", style="rounded,filled", fillcolor="#F2F4FF", color="#8A8FBF", fontname="NanumGothic", fontsize="24")
    dot.attr("edge", color="#666666", fontname="NanumGothic", fontsize="14", arrowsize="0.8")

    emitted: set[str] = set()
    for i, layer in enumerate(layers, start=1):
        layer_name = first_non_empty(layer.get("name"), f"Layer {i}")
        with dot.subgraph(name=f"cluster_{safe_node_id(layer_name)}") as sub:
            sub.attr(label=layer_name, style="rounded,filled", color="#C9C870", fillcolor="#FFFDEB", fontname="NanumGothic", fontsize="22")
            for raw_id in layer.get("component_ids", []):
                sid = raw_to_safe.get(str(raw_id), safe_node_id(raw_id))
                comp = by_safe.get(sid, {"name": raw_id})
                sub.node(sid, component_label(comp))
                emitted.add(sid)

    for comp in components:
        sid = safe_node_id(comp.get("component_id"))
        if sid not in emitted:
            dot.node(sid, component_label(comp))
            emitted.add(sid)

    for rel in select_diagram_relations(relations, components=components, max_edges=14):
        src_raw = str(rel.get("source") or rel.get("from") or rel.get("source_component_id") or "")
        tgt_raw = str(rel.get("target") or rel.get("to") or rel.get("target_component_id") or "")
        src = raw_to_safe.get(src_raw, safe_node_id(src_raw))
        tgt = raw_to_safe.get(tgt_raw, safe_node_id(tgt_raw))
        if src in emitted and tgt in emitted and src != tgt:
            # fallback PNG에서도 edge label은 제거합니다.
            dot.edge(src, tgt)

    out_stem = str(png_path.with_suffix(""))
    rendered = Path(dot.render(out_stem, cleanup=True))
    if rendered != png_path:
        shutil.move(str(rendered), str(png_path))
    return png_path


def clear_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.clear()
    # python-docx는 셀에 최소 1개 paragraph가 있어야 하므로 첫 paragraph를 재사용
    if not cell.paragraphs:
        cell.add_paragraph()


def set_cell_text(cell, text: Any, font_size: int = 10, bold: bool = False) -> None:
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text or ""))
    run.font.size = Pt(font_size)
    run.bold = bold


def fill_meta_table(table: Table, args: argparse.Namespace) -> None:
    # 템플릿 기준: 0번 표 = 제목/시스템명/단계/작성일자/버전
    if len(table.rows) < 3:
        return
    # 병합 셀 구조 때문에 안전하게 label 위치 기준으로 채움
    cells = table.rows[1].cells
    if len(cells) >= 5:
        set_cell_text(cells[1], args.system_name)
        set_cell_text(cells[4], args.subsystem_name)
    cells = table.rows[2].cells
    if len(cells) >= 7:
        set_cell_text(cells[1], "설계")
        set_cell_text(cells[4], args.write_date)
        set_cell_text(cells[6], args.version)


def insert_picture_in_arch_table(table: Table, png_path: Path) -> None:
    cell = table.cell(0, 0)
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    # A4 landscape, 좌우 여백 포함 양식 내부 폭에 맞춤
    run.add_picture(str(png_path), width=Inches(9.3))


def fill_requirement_table(table: Table, item: dict[str, Any]) -> None:
    # 템플릿 기준: 0행 label, 1행 요구사항 내용, 2행 label, 3행 구현방안
    if len(table.rows) < 4:
        return
    set_cell_text(table.rows[0].cells[0], "요구사항 내용", bold=True)
    set_cell_text(table.rows[1].cells[0], clean_requirement_text(item.get("description")), font_size=10)
    set_cell_text(table.rows[2].cells[0], "구현방안", bold=True)
    set_cell_text(table.rows[3].cells[0], clean_requirement_text(item.get("implementation")), font_size=10)


def remove_table(table: Table) -> None:
    tbl = table._tbl
    tbl.getparent().remove(tbl)


def clone_table_after(table: Table) -> Table:
    new_tbl = copy.deepcopy(table._tbl)
    # 표 사이 빈 문단 추가
    spacer = OxmlElement("w:p")
    table._tbl.addnext(spacer)
    spacer.addnext(new_tbl)
    return Table(new_tbl, table._parent)


def keep_together_table(table: Table) -> None:
    """요구사항 1개 표가 가능하면 중간에서 잘리지 않도록 문단 속성을 보강합니다."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                ppr = paragraph._p.get_or_add_pPr()
                keep = OxmlElement("w:keepNext")
                ppr.append(keep)


def render_docx(document_data: dict[str, Any], structure_data: dict[str, Any] | None, args: argparse.Namespace) -> Path:
    template_path = Path(args.template)
    if not template_path.exists():
        raise FileNotFoundError(f"템플릿 파일이 없습니다: {template_path}")

    # 다이어그램은 structure가 있으면 structure 기준, 없으면 document 기준
    diagram_source = structure_data or document_data
    out_path = ensure_parent(args.out)
    diagram_path = ensure_parent(args.diagram_out or out_path.with_name(out_path.stem + "_diagram.png"))
    mermaid_path = ensure_parent(args.mermaid_out or out_path.with_name(out_path.stem + ".mmd"))
    render_diagram(diagram_source, mermaid_path, diagram_path)

    doc = Document(str(template_path))
    if len(doc.tables) < 3:
        raise ValueError("ARCH 템플릿에는 최소 3개의 표(메타, 아키텍처 그림, 요구사항/구현방안)가 필요합니다.")

    fill_meta_table(doc.tables[0], args)
    insert_picture_in_arch_table(doc.tables[1], diagram_path)

    implementations = [i for i in document_data.get("requirement_implementations", []) if isinstance(i, dict)]
    req_table = doc.tables[2]
    if not implementations:
        remove_table(req_table)
    else:
        current = req_table
        fill_requirement_table(current, implementations[0])
        keep_together_table(current)
        for item in implementations[1:]:
            current = clone_table_after(current)
            fill_requirement_table(current, item)
            keep_together_table(current)

    doc.save(str(out_path))
    return out_path


def main() -> None:
    args = parse_args()
    document_data = load_json(args.document)
    structure_data = load_json(args.structure) if args.structure else None
    out = render_docx(document_data, structure_data, args)
    print(f"docx 생성 → {out}")


if __name__ == "__main__":
    main()
