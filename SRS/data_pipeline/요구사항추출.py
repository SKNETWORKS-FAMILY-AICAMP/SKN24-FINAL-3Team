import os
import re
import docx
from pydantic import BaseModel, Field
from typing import List, Optional

# --- 1. 스키마 ---
class CanonicalRequirement(BaseModel):
    requirement_id: str
    requirement_name: str
    requirement_type: str
    description: str
    source: List[str]
    constraints: List[str] = []
    priority: str = "중"
    validation_criteria: List[str] = ["검토 필요"]
    note: Optional[str] = None

class RequirementDocument(BaseModel):
    requirements: List[CanonicalRequirement]

# --- 2. 최적화된 추출 함수 ---
def extract_from_file(file_path):
    doc = docx.Document(file_path)
    requirements_map = {} 
    last_id = None
    
    # 제외할 헤더 키워드들
    exclude_keywords = ["요구사항 명칭", "요구사항 명", "산출정보", "정의", "세부내용", "상세설명", "요구사항"]

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells if cell.text.strip()]
            if not cells: continue
            
            # 1. ID 발견: 발견 즉시 해당 행은 건너뛰고 다음 행부터 데이터 수집
            id_match = next((c for c in cells if re.search(r"[A-Za-z]{3}-\d+", c)), None)
            if id_match:
                last_id = id_match
                if last_id not in requirements_map:
                    requirements_map[last_id] = {"id": last_id, "name": "", "desc": "", "source": os.path.basename(file_path)}
                continue 

            # 2. 데이터 수집
            if last_id:
                # 불필요한 키워드 제거
                clean_parts = [c for c in cells if not any(k in c for k in exclude_keywords)]
                if not clean_parts: continue
                
                # 첫 번째 행은 명칭으로, 나머지는 설명으로 누적
                if not requirements_map[last_id]["name"]:
                    requirements_map[last_id]["name"] = clean_parts[0]
                requirements_map[last_id]["desc"] += " ".join(clean_parts) + " "

    # 3. 객체 생성
    final_reqs = []
    # 제거할 키워드 정의
    blacklist = ["합 계", "ID", "상세 설명", "세부 내용", "구분", "고유번호"]
    
    for data in requirements_map.values():
        name = data["name"].strip()
        desc = data["desc"].strip()
        req_id = data["id"]
        
        # 1. ID가 '000'으로 끝나는 경우 제거 (구조적 노이즈)
        if req_id.endswith("-000"):
            continue
            
        # 2. 명칭이나 설명이 블랙리스트 키워드와 정확히 일치하거나 단순 반복인 경우 제거
        if name in blacklist or desc in blacklist:
            continue
            
        # 3. 설명이 너무 짧은 경우(예: 5자 미만) 제거
        if len(desc) < 10:
            continue
            
        # 4. 필수 필드 검증 후 추가
        final_reqs.append(CanonicalRequirement(
            requirement_id=req_id,
            requirement_name=name,
            requirement_type="비기능" if any(k in (name+desc) for k in ["성능", "보안", "가용성", "호환성"]) else "기능",
            description=desc,
            source=[data["source"]]
        ))
    return final_reqs

# --- 3. 실행부 ---
if __name__ == "__main__":
    folder_path = r"C:\skn24\수업자료\08_large_language_model\00.final_project\data_pipeline\data\RFP"
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            print(f"처리 중: {filename}")
            try:
                reqs = extract_from_file(os.path.join(folder_path, filename))
                result = RequirementDocument(requirements=reqs)
                with open(f"{os.path.splitext(filename)[0]}_final.json", "w", encoding="utf-8") as f:
                    f.write(result.model_dump_json(indent=4))
                print(f"완료: {len(reqs)}개 추출 성공")
            except Exception as e:
                print(f"오류: {e}")