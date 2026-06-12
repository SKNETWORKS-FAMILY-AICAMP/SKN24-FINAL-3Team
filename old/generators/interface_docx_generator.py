import copy

from agents.interface_agent.config import *
from docx import Document
from docx.table import Table
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from agents.interface_agent.extractors import ensure_docx_output_path
from agents.interface_agent.analysis import normalize_ui_structure_data
from generators.interface_image_markers import align_button_markers_to_process_contents

INTERFACE_TEMPLATE_PATH = os.getenv("INTERFACE_TEMPLATE_PATH", "./template/사용자 인터페이스 설계서.docx")

def set_cell_shading(cell, fill: str):
    """DOCX 표 셀에 배경색을 적용합니다."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text_value: str, bold: bool = False, size: int = 9):
    """DOCX 표 셀에 텍스트와 기본 서식을 적용합니다."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text_value) if text_value is not None else "")
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    """DOCX 표 전체에 회색 실선 테두리를 적용합니다."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '808080')
        borders.append(border)
    tblPr.append(borders)


def resolve_interface_template_path(template_path: str = INTERFACE_TEMPLATE_PATH) -> Optional[Path]:
    path = Path(template_path)
    if path.exists():
        return path

    template_dir = Path("./template")
    if template_dir.exists():
        for candidate in template_dir.glob("*.docx"):
            if "사용자 인터페이스" in candidate.name or "interface" in candidate.name.lower():
                return candidate
    return None


def clone_table_after_block(block, table: Table) -> Table:
    new_tbl = copy.deepcopy(table._tbl)
    block.addnext(new_tbl)
    return Table(new_tbl, table._parent)


def insert_paragraph_after_block(doc: Document, block, text_value: str = "", page_break: bool = False):
    paragraph = doc.add_paragraph()
    paragraph._p.getparent().remove(paragraph._p)
    block.addnext(paragraph._p)
    if page_break:
        paragraph.add_run().add_break(WD_BREAK.PAGE)
    if text_value:
        run = paragraph.add_run(text_value)
        run.bold = True
        run.font.size = Pt(10)
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def find_paragraph(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    return None


def fill_repeating_table(table: Table, rows: List[List[str]]):
    if not rows:
        return
    for row_idx, values in enumerate(rows):
        row = table.rows[1] if row_idx == 0 and len(table.rows) > 1 else table.add_row()
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value, size=8)


def fill_template_header_table(doc: Document, requirement_summary: Dict[str, Any]):
    table = doc.tables[0]
    system_name = requirement_summary.get("system_name") or requirement_summary.get("project_name") or ""
    subsystem_name = requirement_summary.get("subsystem_name") or ""
    created_date = datetime.now().strftime("%Y-%m-%d")

    set_cell_text(table.cell(1, 1), system_name, size=8)
    set_cell_text(table.cell(1, 4), subsystem_name, size=8)
    set_cell_text(table.cell(2, 1), "설계", size=8)
    set_cell_text(table.cell(2, 3), created_date, size=8)
    if len(table.rows[2].cells) >= 6:
        set_cell_text(table.cell(2, 5), requirement_summary.get("version", ""), size=8)


def fill_template_structure_table(doc: Document, ui_structure: List[Dict[str, str]]):
    rows = [
        [item.get("level1", ""), item.get("level2", ""), item.get("level3", ""), item.get("level4", "")]
        for item in ui_structure
    ]
    fill_repeating_table(doc.tables[1], rows)


def fill_template_screen_list_table(doc: Document, screen_specs: List[Dict[str, Any]]):
    rows = [[screen.get("screen_id", ""), screen.get("screen_name", "")] for screen in screen_specs]
    fill_repeating_table(doc.tables[2], rows)


def fill_template_screen_detail_table(table: Table, screen_spec: Dict[str, Any]):
    rows = [
        ["화면ID", screen_spec.get("screen_id", ""), "화면명", screen_spec.get("screen_name", "")],
        ["화면유형", screen_spec.get("screen_type", ""), "메뉴경로", screen_spec.get("menu_path", "")],
    ]
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            set_cell_text(table.cell(row_idx, col_idx), value, bold=(col_idx in [0, 2]), size=8)
            if col_idx in [0, 2]:
                set_cell_shading(table.cell(row_idx, col_idx), "D9D9D9")

    set_cell_text(table.cell(2, 0), "화면개요", bold=True, size=8)
    set_cell_shading(table.cell(2, 0), "D9D9D9")
    set_cell_text(table.cell(2, 1), screen_spec.get("screen_overview", ""), size=8)

    image_cell = table.cell(3, 0)
    image_cell.text = ""
    paragraph = image_cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_path = Path(screen_spec.get("annotated_image_path") or screen_spec.get("image_path", ""))
    if image_path.exists():
        try:
            paragraph.add_run().add_picture(str(image_path), width=Inches(6.7))
        except Exception:
            paragraph.add_run(f"[이미지 삽입 실패: {image_path.name}]")
    elif str(image_path):
        paragraph.add_run(f"[이미지 없음: {image_path}]")


def fill_template_process_table(table: Table, process_contents: List[Dict[str, Any]]):
    set_cell_text(table.cell(0, 0), "처리 내용", bold=True, size=9)
    body_cell = table.cell(1, 0)
    body_cell.text = ""
    if not process_contents:
        body_cell.paragraphs[0].add_run("- 처리 내용 없음")
        return

    first = True
    for idx, item in enumerate(process_contents, start=1):
        marker_no = item.get("no", idx)
        title = item.get("title") or f"처리 {marker_no}"
        description = item.get("description", "")
        basis = item.get("requirement_basis", "")

        paragraph = body_cell.paragraphs[0] if first else body_cell.add_paragraph()
        first = False
        paragraph.add_run(f"- {marker_no}. {title}").bold = True
        paragraph.paragraph_format.space_after = Pt(2)

        if description:
            desc_p = body_cell.add_paragraph()
            desc_p.paragraph_format.left_indent = Cm(0.5)
            desc_p.add_run(f"· [{marker_no}] {description}")
        if basis:
            basis_p = body_cell.add_paragraph()
            basis_p.paragraph_format.left_indent = Cm(0.5)
            basis_p.add_run(f"· 근거: {basis}")


def generate_ui_design_docx_from_template(
    requirement_summary: Dict[str, Any],
    ui_structure: Union[List[Dict[str, str]], Dict[str, str]],
    screen_specs: List[Dict[str, Any]],
    output_path: Path,
) -> Path:
    template_path = resolve_interface_template_path()
    if not template_path:
        raise FileNotFoundError(f"사용자 인터페이스 설계서 템플릿을 찾지 못했습니다: {INTERFACE_TEMPLATE_PATH}")

    output_path = ensure_docx_output_path(output_path)
    doc = Document(str(template_path))
    if len(doc.tables) < 5:
        raise ValueError("사용자 인터페이스 설계서 템플릿에는 최소 5개의 표가 필요합니다.")

    ui_structure = normalize_ui_structure_data(ui_structure, screen_specs)
    aligned_screens = []
    for screen in screen_specs:
        screen_copy = dict(screen)
        align_func = globals().get("align_button_markers_to_process_contents")
        if callable(align_func):
            screen_copy = align_func(screen_copy)
        aligned_screens.append(screen_copy)

    fill_template_header_table(doc, requirement_summary)
    fill_template_structure_table(doc, ui_structure)
    fill_template_screen_list_table(doc, aligned_screens)

    detail_template = doc.tables[3]
    process_template = doc.tables[4]
    heading = find_paragraph(doc, "3.1")

    if not aligned_screens:
        fill_template_screen_detail_table(detail_template, {})
        fill_template_process_table(process_template, [])
    else:
        for idx, screen in enumerate(aligned_screens, start=1):
            no_text = str(screen.get("screen_id", f"UI-{idx:03d}")).split("-")[-1]
            no = int(no_text) if no_text.isdigit() else idx
            heading_text = f"3.{no} {screen.get('screen_name', '')}".strip()

            if idx == 1:
                if heading is not None:
                    heading.text = ""
                    run = heading.add_run(heading_text)
                    run.bold = True
                    run.font.size = Pt(10)
                detail_table = detail_template
                process_table = process_template
            else:
                anchor = process_template._tbl if idx == 2 else process_table._tbl
                page_break = insert_paragraph_after_block(doc, anchor, page_break=True)
                heading = insert_paragraph_after_block(doc, page_break._p, heading_text)
                detail_table = clone_table_after_block(heading._p, detail_template)
                blank = insert_paragraph_after_block(doc, detail_table._tbl)
                process_table = clone_table_after_block(blank._p, process_template)

            fill_template_screen_detail_table(detail_table, screen)
            fill_template_process_table(process_table, screen.get("process_contents", []))

    try:
        doc.save(str(output_path))
    except PermissionError:
        alt_path = output_path.with_name(f"{output_path.stem}_new{output_path.suffix}")
        print(f"기존 DOCX가 열려 있어 {alt_path.name} 파일로 저장합니다.")
        doc.save(str(alt_path))
        output_path = alt_path
    return output_path


def add_heading(doc: Document, text_value: str, level: int = 1):
    """DOCX 문서에 설계서용 제목 문단을 추가합니다."""
    p = doc.add_paragraph()
    run = p.add_run(text_value)
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 10)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_simple_table(doc: Document, headers: List[str], rows: List[List[str]], widths: Optional[List[float]] = None):
    """헤더와 행 데이터를 받아 기본 스타일의 DOCX 표를 추가합니다."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "D9D9D9")
        set_cell_text(hdr[i], h, bold=True, size=8)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8)

    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)

    return table


def add_key_value_table(doc: Document, pairs: List[List[str]]):
    """화면 기본정보처럼 키-값 쌍으로 구성된 DOCX 표를 추가합니다."""
    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    for row_pair in pairs:
        row = table.add_row().cells
        for i, value in enumerate(row_pair):
            set_cell_text(row[i], value, bold=(i in [0,2]), size=8)
            if i in [0,2]:
                set_cell_shading(row[i], "D9D9D9")
    return table



def add_process_contents_box(doc: Document, process_contents: List[Dict[str, Any]]):
    """화면 상세 설계 하단에 예시 양식과 같은 처리 내용 영역을 추가합니다."""
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_cell_shading(table.cell(0, 0), "D9D9D9")
    set_cell_text(table.cell(0, 0), "처리 내용", bold=True, size=9)

    body_cell = table.cell(1, 0)
    body_cell.text = ""
    if not process_contents:
        body_cell.paragraphs[0].add_run("- 처리 내용 없음")
        return table

    first = True
    for idx, item in enumerate(process_contents, start=1):
        marker_no = item.get("no", idx)
        title = item.get("title") or f"처리 {marker_no}"
        description = item.get("description", "")
        basis = item.get("requirement_basis", "")

        p = body_cell.paragraphs[0] if first else body_cell.add_paragraph()
        first = False
        p.add_run(f"- {marker_no}. {title}").bold = True
        p.paragraph_format.space_after = Pt(2)

        if description:
            desc_p = body_cell.add_paragraph()
            desc_p.paragraph_format.left_indent = Cm(0.5)
            desc_p.add_run(f"· [{marker_no}] {description}")
        if basis:
            basis_p = body_cell.add_paragraph()
            basis_p.paragraph_format.left_indent = Cm(0.5)
            basis_p.add_run(f"· 근거: {basis}")
    return table

def add_screen_detail_table_with_image(doc: Document, screen_spec: Dict[str, Any]):
    """화면 기본정보와 번호 버튼 이미지를 하나의 상세 설계 표 안에 넣습니다."""
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    rows = [
        ["화면ID", screen_spec.get("screen_id", ""), "화면명", screen_spec.get("screen_name", "")],
        ["화면유형", screen_spec.get("screen_type", ""), "메뉴경로", screen_spec.get("menu_path", "")],
        ["화면개요", screen_spec.get("screen_overview", ""), "", ""],
    ]
    for r_i, row_data in enumerate(rows):
        for c_i, value in enumerate(row_data):
            set_cell_text(table.cell(r_i, c_i), value, bold=(c_i in [0, 2]), size=8)
            if c_i in [0, 2]:
                set_cell_shading(table.cell(r_i, c_i), "D9D9D9")

    table.cell(2, 1).merge(table.cell(2, 3))
    image_cell = table.cell(3, 0).merge(table.cell(3, 3))
    image_cell.text = ""
    image_path = Path(screen_spec.get("annotated_image_path") or screen_spec.get("image_path", ""))
    p = image_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path.exists():
        try:
            p.add_run().add_picture(str(image_path), width=Inches(6.7))
        except Exception:
            p.add_run(f"[이미지 삽입 실패: {image_path.name}]")
    else:
        p.add_run(f"[이미지 없음: {image_path}]")
    return table



def build_integrated_ui_design_data(
    requirement_summary: Dict[str, Any],
    ui_structure: Union[List[Dict[str, str]], Dict[str, str]],
    screen_specs: List[Dict[str, Any]],
    output_docx_path: Path
) -> Dict[str, Any]:
    """화면설계서 생성에 필요한 데이터를 하나의 통합 JSON 구조로 조립합니다."""
    ui_structure = normalize_ui_structure_data(ui_structure, screen_specs)
    align_func = globals().get("align_button_markers_to_process_contents")
    screen_details = []
    for screen in screen_specs:
        screen_copy = dict(screen)
        if callable(align_func):
            screen_copy = align_func(screen_copy)
        screen_details.append(screen_copy)
    created_date = datetime.now().strftime("%Y-%m-%d")
    return {
        "document_header": {
            "title": "사용자 인터페이스 설계서",
            "system_name": requirement_summary.get("system_name") or requirement_summary.get("project_name") or "",
            "subsystem_name": requirement_summary.get("subsystem_name") or "",
            "phase": "설계",
            "created_date": created_date,
            "version": "",
        },
        "ui_structure": ui_structure,
        "ui_screen_list": [
            {
                "screen_id": screen.get("screen_id", ""),
                "screen_name": screen.get("screen_name", ""),
            }
            for screen in screen_details
        ],
        "screen_details": screen_details,
        "source_requirements": requirement_summary,
        "outputs": {
            "docx_path": str(output_docx_path),
            "generated_at": datetime.now().isoformat(timespec="seconds"),
        },
    }


def save_integrated_ui_design_json(
    requirement_summary: Dict[str, Any],
    ui_structure: Union[List[Dict[str, str]], Dict[str, str]],
    screen_specs: List[Dict[str, Any]],
    output_docx_path: Path,
    output_json_path: Path
) -> Path:
    """통합 화면설계서 JSON을 저장합니다."""
    output_json_path = Path(output_json_path)
    output_json_path.parent.mkdir(parents=True, exist_ok=True)
    integrated_data = build_integrated_ui_design_data(
        requirement_summary=requirement_summary,
        ui_structure=ui_structure,
        screen_specs=screen_specs,
        output_docx_path=output_docx_path,
    )
    output_json_path.write_text(
        json.dumps(integrated_data, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return output_json_path

def create_ui_design_docx(
    requirement_summary: Dict[str, Any],
    ui_structure: Union[List[Dict[str, str]], Dict[str, str]],
    screen_specs: List[Dict[str, Any]],
    output_path: Path
):
    """확정된 화면설계서 양식 순서에 맞춰 DOCX 파일을 생성합니다."""
    output_path = ensure_docx_output_path(output_path)
    try:
        return generate_ui_design_docx_from_template(
            requirement_summary=requirement_summary,
            ui_structure=ui_structure,
            screen_specs=screen_specs,
            output_path=output_path,
        )
    except Exception as exc:
        print(f"[WARN] UI 설계서 템플릿 기반 DOCX 생성 실패, 기존 생성 방식으로 전환합니다: {type(exc).__name__}: {exc}")

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    styles = doc.styles
    styles['Normal'].font.name = '맑은 고딕'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    styles['Normal'].font.size = Pt(9)

    title_table = doc.add_table(rows=3, cols=5)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(title_table)

    title_table.cell(0, 0).merge(title_table.cell(0, 4))
    set_cell_text(title_table.cell(0, 0), "사용자 인터페이스 설계서", bold=True, size=14)
    title_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    system_name = requirement_summary.get("system_name") or requirement_summary.get("project_name") or ""
    subsystem_name = requirement_summary.get("subsystem_name") or ""

    header_rows = [
        ["시스템명", system_name, "서브시스템명", subsystem_name, ""],
        ["단계명", "설계", "작성일자", datetime.now().strftime("%Y-%m-%d"), "버전"],
    ]
    for r_i, row_data in enumerate(header_rows, start=1):
        for c_i, value in enumerate(row_data):
            set_cell_text(title_table.cell(r_i, c_i), value, bold=(c_i in [0, 2, 4]), size=8)
            if c_i in [0, 2, 4]:
                set_cell_shading(title_table.cell(r_i, c_i), "D9D9D9")

    doc.add_paragraph("")

    add_heading(doc, "1. 사용자 인터페이스 구조도", level=1)
    ui_structure = normalize_ui_structure_data(ui_structure, screen_specs)
    rows = [[item.get("level1", ""), item.get("level2", ""), item.get("level3", ""), item.get("level4", "")] for item in ui_structure]
    add_simple_table(doc, ["업무 영역\n(Level 1)", "Level 2", "Level 3", "Level 4"], rows, widths=[3.0, 3.0, 3.5, 5.5])

    add_heading(doc, "2. 사용자 인터페이스 목록", level=1)
    rows = [[s.get("screen_id", ""), s.get("screen_name", "")] for s in screen_specs]
    add_simple_table(doc, ["화면 ID", "화면명"], rows, widths=[4.0, 10.0])

    add_heading(doc, "3. 화면 상세 설계", level=1)

    for s in screen_specs:
        align_func = globals().get("align_button_markers_to_process_contents")
        if callable(align_func):
            s = align_func(s)
        doc.add_page_break()
        no_text = str(s.get("screen_id", "UI-000")).split("-")[-1]
        no = int(no_text) if no_text.isdigit() else len(doc.paragraphs)
        add_heading(doc, f"3.{no} {s.get('screen_name', '')}", level=2)
        add_screen_detail_table_with_image(doc, s)
        doc.add_paragraph("")
        add_process_contents_box(doc, s.get("process_contents", []))

    try:
        doc.save(str(output_path))
    except PermissionError:
        alt_path = output_path.with_name(f"{output_path.stem}_new{output_path.suffix}")
        print(f"기존 DOCX가 열려 있어 {alt_path.name} 파일로 저장합니다.")
        doc.save(str(alt_path))
        output_path = alt_path
    return output_path
