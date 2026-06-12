import os
from datetime import date
from pathlib import Path
from typing import Any, Dict

from docx import Document
from generators.common.docx_utils import clean_text, clone_table_after, save_docx_with_fallback, set_cell

try:
    from dotenv import load_dotenv
except ImportError:
    def load_dotenv():
        return None

load_dotenv()

TEMPLATE_PATH = os.getenv("DB_DESIGN_TEMPLATE_PATH", "./template/데이터베이스 설계서.docx")
OUTPUT_PATH = os.getenv("DB_DESIGN_OUTPUT_PATH", f"./output/데이터베이스 설계서_{date.today()}.docx")


def fill_header_table(doc, design: Dict[str, Any]):
    table = doc.tables[0]
    set_cell(table.cell(1, 1), design.get("system_name", ""))
    set_cell(table.cell(1, 4), design.get("subsystem_name", ""))
    set_cell(table.cell(2, 1), design.get("stage_name", "설계"))
    set_cell(table.cell(2, 4), design.get("created_date", str(date.today())))
    set_cell(table.cell(2, 6), design.get("version", "v1.0"))


def fill_database_list_table(doc, design: Dict[str, Any]):
    table = doc.tables[1]
    databases = design.get("databases", [])
    if not databases:
        return

    for idx, database in enumerate(databases):
        row = table.rows[2] if idx == 0 else table.add_row()
        values = [
            database.get("database_id", ""),
            database.get("database_name", ""),
            database.get("owner_department", ""),
            database.get("note", ""),
            "",
        ]
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)


def fill_database_definition_table(doc, design: Dict[str, Any]):
    table = doc.tables[2]
    databases = design.get("databases", [])
    if not databases:
        return

    database = databases[0]
    set_cell(table.cell(0, 1), database.get("database_id", ""))
    set_cell(table.cell(0, 3), database.get("database_name", ""))
    set_cell(table.cell(0, 5), database.get("storage_group", ""))
    set_cell(table.cell(1, 1), database.get("bufferpool", ""))
    set_cell(table.cell(1, 5), database.get("index_bufferpool", ""))

    for idx, item in enumerate(design.get("tables", [])):
        row = table.rows[3] if idx == 0 else table.add_row()
        values = [
            item.get("tablespace_name", ""),
            item.get("capacity", ""),
            item.get("table_id", ""),
            item.get("table_name", ""),
            f"IX_{item.get('table_name', '')}"[:30],
            "산정 필요",
            item.get("note", ""),
        ]
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)


def fill_table_spec(table, item: Dict[str, Any]):
    set_cell(table.cell(0, 1), item.get("table_id", ""))
    set_cell(table.cell(0, 5), item.get("table_name", ""))
    set_cell(table.cell(1, 1), item.get("database_name", ""))
    set_cell(table.cell(1, 5), item.get("tablespace_name", ""))
    set_cell(table.cell(2, 1), item.get("trigger_config", ""))
    set_cell(table.cell(3, 1), item.get("table_description", ""))
    set_cell(table.cell(5, 0), item.get("initial_count", ""))
    set_cell(table.cell(5, 1), item.get("daily_growth", ""))
    set_cell(table.cell(5, 2), item.get("retention_period", ""))
    set_cell(table.cell(5, 3), item.get("max_count", ""))
    set_cell(table.cell(5, 4), item.get("capacity", ""))
    set_cell(table.cell(5, 5), item.get("note", ""))

    for idx, column in enumerate(item.get("columns", [])):
        row = table.rows[7] if idx == 0 else table.add_row()
        values = [
            column.get("column_name", ""),
            column.get("column_id", ""),
            column.get("type_and_length", ""),
            column.get("not_null", ""),
            column.get("pk", ""),
            column.get("fk", ""),
            column.get("idx", ""),
            column.get("default", ""),
            column.get("constraint", ""),
        ]
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)


def fill_table_spec_tables(doc, design: Dict[str, Any]):
    tables = design.get("tables", [])
    if not tables:
        return

    template_table = doc.tables[3]
    for _ in tables[1:]:
        clone_table_after(template_table)

    for table, item in zip(doc.tables[3:3 + len(tables)], tables):
        fill_table_spec(table, item)


def generate_database_design_docx(
    design: Dict[str, Any],
    template_path: str = TEMPLATE_PATH,
    output_path: str = OUTPUT_PATH,
) -> str:
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc = Document(template_path)

    fill_header_table(doc, design)
    fill_database_list_table(doc, design)
    fill_database_definition_table(doc, design)
    fill_table_spec_tables(doc, design)

    saved_path = save_docx_with_fallback(doc, output_path)
    if saved_path != output_path:
        print(f"[WARN] 기존 파일을 덮어쓸 수 없어 다른 이름으로 저장했습니다: {saved_path}")
    print(f"[완료] 데이터베이스 설계서: {saved_path}")
    return saved_path


if __name__ == "__main__":
    import json

    with open("./json_temp/database_design_agent_output.json", "r", encoding="utf-8") as f:
        design = json.load(f)

    generate_database_design_docx(design)
