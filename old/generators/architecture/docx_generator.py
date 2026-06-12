import os
from datetime import date
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.shared import Inches
from generators.common.docx_utils import clean_text, clone_table_after, save_docx_with_fallback, set_cell

try:
    from dotenv import load_dotenv
except ImportError:
    def load_dotenv():
        return None

load_dotenv()

TEMPLATE_PATH = os.getenv("ARCHITECTURE_TEMPLATE_PATH", "./template/아키텍처 설계서.docx")
OUTPUT_PATH = os.getenv("ARCHITECTURE_OUTPUT_PATH", f"./output/아키텍처 설계서_{date.today()}.docx")


def resolve_template_path(template_path: str = TEMPLATE_PATH) -> str:
    path = Path(template_path)
    if path.exists():
        return str(path)

    template_dir = Path("./template")
    if template_dir.exists():
        for candidate in template_dir.glob("*.docx"):
            if "아키텍처" in candidate.name or "architecture" in candidate.name.lower():
                return str(candidate)

    raise FileNotFoundError(f"아키텍처 설계서 템플릿을 찾지 못했습니다: {template_path}")


def add_cell_paragraph(cell, value: Any, *, bold: bool = False) -> None:
    text = clean_text(value)
    if not text:
        return
    paragraph = cell.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold


def fill_header_table(doc, payload: Dict[str, Any]) -> None:
    table = doc.tables[0]
    infra = payload.get("user_infra_spec", {})

    system_name = payload.get("system_name") or "온프레미스 생성형 AI 플랫폼"
    subsystem_name = payload.get("subsystem_name") or "아키텍처"

    set_cell(table.cell(1, 1), system_name)
    set_cell(table.cell(1, 4), subsystem_name)
    set_cell(table.cell(2, 1), payload.get("stage_name", "설계"))
    set_cell(table.cell(2, 3), payload.get("created_date", str(date.today())))
    set_cell(table.cell(2, 5), payload.get("version", "v1.0"))

    if infra.get("system_name"):
        set_cell(table.cell(1, 1), infra.get("system_name"))


def build_system_architecture_text(payload: Dict[str, Any]) -> List[str]:
    infra = payload.get("user_infra_spec", {})
    extracted = payload.get("extracted_infra", {})

    lines = [
        "[시스템 구성 개요]",
        clean_text(extracted.get("system_architecture")),
        "",
        "[시스템 소프트웨어 및 미들웨어]",
        clean_text(extracted.get("selected_middleware")),
        "",
        "[하드웨어/네트워크/운영 환경]",
        f"구축 형태: {'Cloud' if infra.get('is_cloud') else 'On-Premise'}",
        f"미들웨어 스택: {clean_text(infra.get('middleware_stack'))}",
        f"예상 동시 사용자: {clean_text(infra.get('expected_ccu'))}",
        f"서버 사양: {clean_text(infra.get('server_hardware_spec'))}",
        f"방화벽 구성: {clean_text(infra.get('firewall_setting'))}",
        f"인증/보안 연계: {clean_text(infra.get('security_auth'))}",
        "",
        "[보안 아키텍처]",
        clean_text(extracted.get("security_architecture")),
    ]
    return [line for line in lines if line is not None]


def fill_system_architecture_table(doc, payload: Dict[str, Any], image_path: str | None = None) -> None:
    table = doc.tables[1]
    cell = table.cell(0, 0)
    cell.text = ""

    if image_path and Path(image_path).exists():
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(6.2))


def requirement_items(payload: Dict[str, Any]) -> List[Dict[str, Any]]:
    return payload.get("requirement_doc", {}).get("requirements", [])


def analyzed_map(payload: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    result = {}
    for item in payload.get("analyzed_reqs", []):
        req_id = item.get("requirement_id")
        if req_id:
            result[req_id] = item
    return result


def build_implementation_text(req: Dict[str, Any], analysis: Dict[str, Any], payload: Dict[str, Any]) -> str:
    extracted = payload.get("extracted_infra", {})
    parts = []

    needs = clean_text(analysis.get("implied_middleware_needs"))
    constraints = clean_text(analysis.get("technical_constraints"))
    security = clean_text(extracted.get("security_architecture"))

    if needs:
        parts.append(f"관련 구성요소/미들웨어:\n{needs}")
    if constraints:
        parts.append(f"아키텍처 제약 반영:\n{constraints}")
    if security and (
        "보안" in clean_text(req.get("requirement_name"))
        or "SSO" in clean_text(req.get("requirement_name"))
        or "ERP" in clean_text(req.get("requirement_name"))
        or "인증" in clean_text(req.get("description"))
    ):
        parts.append(f"보안 구현방안:\n{security}")

    if not parts:
        parts.append("시스템 아키텍처 구성요소 및 운영 환경에 반영하여 구현한다.")

    return "\n\n".join(parts)


def fill_requirement_table(table, req: Dict[str, Any], analysis: Dict[str, Any], payload: Dict[str, Any]) -> None:
    set_cell(table.cell(0, 1), req.get("requirement_id", ""))
    set_cell(table.cell(2, 0), req.get("description", ""))
    set_cell(table.cell(4, 0), build_implementation_text(req, analysis, payload))


def fill_requirement_tables(doc, payload: Dict[str, Any]) -> None:
    requirements = requirement_items(payload)
    if not requirements:
        return

    template_table = doc.tables[2]
    for _ in requirements[1:]:
        clone_table_after(template_table)

    analysis_by_id = analyzed_map(payload)
    tables = doc.tables[2:2 + len(requirements)]
    for table, req in zip(tables, requirements):
        fill_requirement_table(table, req, analysis_by_id.get(req.get("requirement_id"), {}), payload)


def generate_architecture_docx(
    payload: Dict[str, Any],
    template_path: str = TEMPLATE_PATH,
    output_path: str = OUTPUT_PATH,
) -> str:
    doc = Document(resolve_template_path(template_path))

    fill_header_table(doc, payload)
    fill_system_architecture_table(doc, payload, payload.get("image_path"))
    fill_requirement_tables(doc, payload)

    saved_path = save_docx_with_fallback(doc, output_path)
    if saved_path != output_path:
        print(f"[WARN] 기존 파일을 덮어쓸 수 없어 다른 이름으로 저장했습니다: {saved_path}")
    print(f"[완료] 아키텍처 설계서: {saved_path}")
    return saved_path


if __name__ == "__main__":
    import json

    with open("./json_temp/architecture_agent_output.json", encoding="utf-8") as f:
        data = json.load(f)

    generate_architecture_docx(data)
