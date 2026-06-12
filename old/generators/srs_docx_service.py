import json
import os
import shutil
import subprocess
from pathlib import Path
from datetime import datetime

_JS_PATH  = Path(__file__).parent / "srs_gen_req_docx.js"
_OUT_DIR  = Path(__file__).parent.parent / "output"
_ROOT_DIR = Path(__file__).parent
_PROJECT_ROOT = Path(__file__).parent.parent


def generate_docx(
    reqs: list[dict],
    prefix: str = "requirements",
    output_path: str | None = None,
) -> str:
    _OUT_DIR.mkdir(exist_ok=True)

    ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = Path(output_path) if output_path else _OUT_DIR / f"{prefix}_{ts}.docx"
    if not out_path.is_absolute():
        out_path = _PROJECT_ROOT / out_path
    out_path.parent.mkdir(parents=True, exist_ok=True)

    clean    = [{k: v for k, v in r.items() if not k.startswith("_")} for r in reqs]

    # tmp 파일을 한글 없는 req_agent 루트에 저장
    tmp_json = _ROOT_DIR / f"_tmp_{ts}.json"
    tmp_js   = _ROOT_DIR / f"_tmp_{ts}.js"

    tmp_json.write_text(json.dumps(clean, ensure_ascii=False), encoding="utf-8")

    script = _JS_PATH.read_text(encoding="utf-8")
    script = script.replace(
        "'/home/claude/sample_reqs.json'",
        repr(tmp_json.as_posix())
    ).replace(
        "'/home/claude/requirements_definition.docx'",
        repr(out_path.as_posix())
    )

    tmp_js.write_text(script, encoding="utf-8")

    node_bin = shutil.which("node")
    if not node_bin:
        tmp_json.unlink(missing_ok=True)
        tmp_js.unlink(missing_ok=True)
        print("[WARN] Node.js 실행 파일을 찾지 못해 Python DOCX fallback으로 생성합니다.")
        return _generate_docx_with_python(clean, out_path)

    try:
        env = os.environ.copy()
        node_paths = [
            _PROJECT_ROOT / "node_modules",
            _PROJECT_ROOT / "SRS" / "req_agent" / "node_modules",
        ]
        existing_node_path = env.get("NODE_PATH")
        env["NODE_PATH"] = os.pathsep.join(
            [str(path) for path in node_paths if path.exists()]
            + ([existing_node_path] if existing_node_path else [])
        )
        result = subprocess.run(
            [node_bin, tmp_js.name],   # 파일명만 (cwd 기준)
            check=True,
            capture_output=True,
            text=True,
            encoding="utf-8",
            cwd=str(_ROOT_DIR),
            env=env,
        )
        print(result.stdout)
    except subprocess.CalledProcessError as e:
        print("=== Node.js 에러 ===")
        print(e.stderr)
        print("[WARN] Python DOCX fallback으로 생성합니다.")
        return _generate_docx_with_python(clean, out_path)
    finally:
        tmp_json.unlink(missing_ok=True)
        tmp_js.unlink(missing_ok=True)

    return str(out_path)


def _generate_docx_with_python(reqs: list[dict], out_path: Path) -> str:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    status_colors = {
        "신규": "E8F5E9",
        "수정": "E3F2FD",
        "기존": "FFFFFF",
    }

    def normalize(value):
        if value is None:
            return ""
        if isinstance(value, list):
            return "\n".join(str(item) for item in value)
        return str(value)

    def set_cell_bg(cell, color: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)
        tc_pr.append(shd)

    def add_cell_text(cell, text, *, bold=False, center=False, fill=None) -> None:
        cell.text = ""
        if fill:
            set_cell_bg(cell, fill)
        paragraph = cell.paragraphs[0]
        if center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(normalize(text))
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(8)
        run.bold = bold
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("사용자 요구사항 정의서")
    title_run.bold = True
    title_run.font.name = "Malgun Gothic"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta_run = meta.add_run(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    meta_run.font.name = "Malgun Gothic"
    meta_run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    meta_run.font.size = Pt(8)

    headers = [
        "요구사항 ID",
        "요구사항명",
        "구분",
        "요구사항 설명",
        "요구사항 출처",
        "제약사항",
        "중요도",
        "해결방안",
        "검수기준",
        "상태",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    for cell, header in zip(table.rows[0].cells, headers):
        add_cell_text(cell, header, bold=True, center=True, fill="D9D9D9")

    for req in reqs:
        status = req.get("status") or "기존"
        fill = status_colors.get(status, "FFFFFF")
        values = [
            req.get("requirement_id", ""),
            req.get("requirement_name", ""),
            req.get("requirement_type", ""),
            req.get("description", ""),
            req.get("source", ""),
            req.get("constraints", ""),
            req.get("priority", ""),
            req.get("note", ""),
            req.get("validation_criteria", ""),
            status,
        ]
        row = table.add_row().cells
        for idx, value in enumerate(values):
            add_cell_text(row[idx], value, center=idx in {0, 2, 6, 9}, fill=fill)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return str(out_path)
