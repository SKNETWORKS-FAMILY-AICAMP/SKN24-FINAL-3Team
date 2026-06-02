from pathlib import Path


def read_file_text(file_path: str, file_ext: str | None = None) -> str:
    path = Path(file_path)
    ext = (file_ext or path.suffix.lstrip(".")).lower()

    if not path.exists():
        raise FileNotFoundError(f"파일을 찾지 못했습니다: {file_path}")

    if ext == "txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    if ext == "docx":
        return _read_docx(path)
    if ext == "pdf":
        return _read_pdf(path)
    if ext in {"xlsx", "xls"}:
        return _read_xlsx(path)

    raise ValueError(f"지원하지 않는 파일 확장자입니다: {ext}")


def _read_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _read_pdf(path: Path) -> str:
    import fitz

    text = []
    with fitz.open(str(path)) as document:
        for page in document:
            text.append(page.get_text())
    return "\n".join(text)


def _read_xlsx(path: Path) -> str:
    import pandas as pd

    sheets = pd.read_excel(str(path), sheet_name=None, dtype=str)
    parts = []
    for sheet_name, frame in sheets.items():
        parts.append(f"[Sheet: {sheet_name}]")
        parts.append(frame.fillna("").to_csv(index=False, sep="\t"))
    return "\n".join(parts)
