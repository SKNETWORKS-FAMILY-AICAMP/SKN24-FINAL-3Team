import tempfile
import unittest
import base64
from pathlib import Path

import fitz
from docx import Document

from tools.llm.json_repair import repair_json_text
from tools.llm.llm_client import LLMClient
from tools.llm.response_parser import parse_json_response
from tools.llm.send_api import send_parallel
from tools.parser.docx_parser import parse_docx
from tools.parser.erd_docx_parser import parse_erd_docx
from tools.parser.pdf_parser import parse_pdf
from tools.parser.rfp_rule_parser import parse_rfp_requirements
from tools.parser.table_parser import parse_tables
from tools.search.search_router import search
from tools.search.search_schema import SearchRequest
from tools.storage.cleanup_manager import cleanup_paths
from tools.storage.downloader import download_file
from tools.storage.uploader import upload_file
from tools.vector.embedding_writer import write_non_functional_requirements
from tools.docx.docx_exporter import export_docx
from tools.docx import docx_exporter
import tools.mermaid.mermaid_renderer as mermaid_renderer


class CommonToolsTest(unittest.TestCase):
    def test_erd_docx_parser_extracts_exported_entity_tables(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            docx_path = Path(root) / "erd.docx"
            document = Document()
            table = document.add_table(rows=5, cols=10)
            table.cell(0, 2).text = "ENT-001"
            table.cell(0, 7).text = "사용자"
            table.cell(1, 4).text = "사용자 정보를 관리하는 테이블입니다."
            table.rows[2].cells[0].text = "속성명"
            table.rows[2].cells[1].text = "동의어"
            table.rows[2].cells[2].text = "데이터타입"
            table.rows[3].cells[0].text = "USER_SN"
            table.rows[3].cells[1].text = "사용자 번호"
            table.rows[3].cells[2].text = "BIGINT"
            table.rows[3].cells[4].text = "Y"
            table.rows[3].cells[5].text = "Y"
            table.rows[4].cells[0].text = "USER_NM"
            table.rows[4].cells[1].text = "사용자 명"
            table.rows[4].cells[2].text = "VARCHAR"
            table.rows[4].cells[3].text = "200"
            document.save(docx_path)

            result = parse_erd_docx(str(docx_path))

            self.assertTrue(result["success"])
            tables = result["data"]["tables"]
            self.assertEqual(tables[0]["logical_name"], "사용자")
            self.assertEqual(tables[0]["description"], "사용자 정보를 관리하는 테이블입니다.")
            self.assertEqual(tables[0]["columns"][0]["physical_name"], "USER_SN")
            self.assertEqual(tables[0]["columns"][0]["constraints"], ["PK"])
            self.assertFalse(tables[0]["columns"][0]["nullable"])
            self.assertEqual(tables[0]["columns"][1]["data_type"], "VARCHAR(200)")

    def test_srs_template_maps_task3_fields_to_cbd_columns(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            output_path = Path(root) / "task3-srs.docx"
            result = export_docx(
                {
                    "docs_cd": "SRS",
                    "content": {
                        "requirement_json_list": [
                            {
                                "gold_id": "GOLD-001",
                                "action_type": "산출",
                                "requirement_name": "CXL 메모리 프레임워크",
                                "requirement_detail": "CXL 메모리 프레임워크를 설계한다.",
                                "sources": ["SFR-001", "SFR-003"],
                                "merge_basis": "중복 기능을 통합함.",
                            }
                        ]
                    },
                    "image_paths": [],
                },
                str(output_path),
                template_path="templates/srs_template.docx",
            )

            self.assertTrue(result["success"])
            row = Document(output_path).tables[1].rows[1]
            self.assertEqual(row.cells[0].text, "GOLD-001")
            self.assertEqual(row.cells[2].text, "산출")
            self.assertEqual(row.cells[3].text, "CXL 메모리 프레임워크를 설계한다.")
            self.assertEqual(row.cells[4].text, "SFR-001\nSFR-003")
            self.assertEqual(row.cells[7].text, "중복 기능을 통합함.")

    def test_db_template_uses_logical_and_physical_column_fields(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            output_path = Path(root) / "db.docx"
            result = export_docx(
                {
                    "docs_cd": "DB",
                    "content": {
                        "db_design_json": {
                            "database_id": "DB-001",
                            "database_name": "업무 DB",
                            "tables": [
                                {
                                    "table_id": "tbl_user",
                                    "table_name": "tbl_user",
                                    "table_logical_name": "사용자",
                                    "database_name": "업무 DB",
                                    "tablespace_name": "TS_USER",
                                    "trigger_config": "해당 없음",
                                    "table_description": "사용자 정보를 관리하는 테이블입니다.",
                                    "columns": [
                                        {
                                            "column_name": "user_sn",
                                            "column_id": "user_sn",
                                            "column_logical_name": "사용자 번호",
                                            "type_and_length": "BIGINT",
                                            "nullable": False,
                                            "pk": "Y",
                                            "fk": "",
                                            "idx": "Y",
                                            "default": "",
                                            "constraint": "",
                                        }
                                    ],
                                }
                            ],
                        }
                    },
                },
                str(output_path),
                template_path="templates/db_template.docx",
            )

            self.assertTrue(result["success"])
            spec = Document(output_path).tables[3]
            self.assertEqual(spec.cell(0, 1).text, "tbl_user")
            self.assertEqual(spec.cell(0, 5).text, "테이블명")
            self.assertEqual(spec.cell(0, 6).text, "사용자")
            self.assertEqual(spec.cell(1, 5).text, "TS명")
            self.assertEqual(spec.cell(1, 6).text, "TS_USER")
            self.assertEqual(spec.cell(5, 0).text, "0")
            self.assertEqual(spec.cell(5, 1).text, "산정 필요")
            self.assertEqual(spec.cell(5, 2).text, "업무 기준에 따름")
            self.assertEqual(spec.cell(5, 3).text, "산정 필요")
            self.assertEqual(spec.cell(5, 4).text, "산정 필요")
            self.assertEqual(spec.cell(7, 0).text, "사용자 번호")
            self.assertEqual(spec.cell(7, 1).text, "user_sn")
            self.assertEqual(spec.cell(7, 8).text, "")

    def test_erd_template_uses_logical_attribute_name_and_split_type_length(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            output_path = Path(root) / "erd.docx"
            result = export_docx(
                {
                    "docs_cd": "ERD",
                    "content": {
                        "erd_entity_json": {
                            "erd_id": "JJ_ERD_040",
                            "erd_name": "거래중계",
                            "tables": [
                                {
                                    "entity_id": "JJ_EN_110",
                                    "logical_name": "기관거래",
                                    "description": "기관거래 정보를 저장합니다.",
                                    "columns": [
                                        {
                                            "logical_name": "거래일자",
                                            "physical_name": "dlng_ymd",
                                            "synonym": "",
                                            "data_type": "CHAR",
                                            "length": "8",
                                            "nullable": False,
                                            "constraints": ["PK", "AUTO_INCREMENT"],
                                            "default": "",
                                        }
                                    ],
                                }
                            ],
                        }
                    },
                },
                str(output_path),
                template_path="templates/erd_template.docx",
            )

            self.assertTrue(result["success"])
            entity = Document(output_path).tables[2]
            self.assertEqual(entity.cell(3, 0).text, "거래일자")
            self.assertEqual(entity.cell(3, 1).text, "")
            self.assertEqual(entity.cell(3, 2).text, "CHAR")
            self.assertEqual(entity.cell(3, 3).text, "8")
            self.assertEqual(entity.cell(3, 4).text, "Y")
            self.assertEqual(entity.cell(3, 5).text, "PK")
            self.assertEqual(entity.cell(3, 9).text, "AUTO_INCREMENT")

    def test_erd_template_entity_name_uses_korean_logical_name(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            output_path = Path(root) / "erd.docx"
            result = export_docx(
                {
                    "docs_cd": "ERD",
                    "content": {
                        "erd_entity_json": {
                            "tables": [
                                {
                                    "entity_id": "ENT-001",
                                    "entity_name": "사용자 그룹",
                                    "logical_name": "tbl_user_group",
                                    "physical_name": "tbl_user_group",
                                    "description": "사용자 그룹 정보를 관리합니다.",
                                    "columns": [],
                                }
                            ],
                        }
                    },
                },
                str(output_path),
                template_path="templates/erd_template.docx",
            )

            self.assertTrue(result["success"])
            entity = Document(output_path).tables[2]
            self.assertEqual(entity.cell(0, 7).text, "사용자 그룹")

    def test_erd_column_to_row_uses_short_intuitive_attribute_name(self) -> None:
        dash_row = docx_exporter._erd_column_to_row(
            {
                "logical_name": "-",
                "physical_name": "rag_ops_nm",
                "data_type": "VARCHAR(100)",
                "nullable": False,
                "constraints": [],
            }
        )
        long_row = docx_exporter._erd_column_to_row(
            {
                "logical_name": "검색증강생성 기본사항 RAG RAGOps 정보를 관리하는 테이블입니다",
                "physical_name": "rag_ops_cn",
                "data_type": "VARCHAR(200)",
                "nullable": False,
                "constraints": [],
            }
        )

        self.assertEqual(dash_row[0], "명")
        self.assertEqual(long_row[0], "내용")

    def test_arch_implementation_text_focuses_on_component_context(self) -> None:
        arch_doc = {
            "components": [
                {
                    "component_id": "WEB",
                    "name": "Web Client",
                    "layer": "Presentation Layer",
                    "description": "사용자 화면과 입력 흐름을 제공",
                },
                {
                    "component_id": "API",
                    "name": "API Server",
                    "layer": "Application Layer",
                    "description": "업무 API 처리",
                },
            ],
            "relations": [
                {"source": "WEB", "target": "API", "description": "HTTP/API 요청"},
                {"source": "API", "target": "RDBMS", "description": "메타데이터 조회"},
            ],
            "deployment_environment": {
                "environment": "운영/검증 분리",
                "web_was": "WEB/WAS 분리",
            },
        }

        text = docx_exporter._arch_implementation_text(arch_doc["components"][0], arch_doc)

        self.assertIn("Web Client는 Presentation Layer에 배치", text)
        self.assertIn("사용자 화면과 입력 흐름을 제공하도록 설계", text)
        self.assertIn("WEB -> API: HTTP/API 요청", text)
        self.assertIn("운영/검증 분리", text)
        self.assertNotIn("구성요소: Web Client, API Server", text)

    def test_arch_template_does_not_require_requirement_id_row(self) -> None:
        table = Document().add_table(rows=4, cols=2)
        table.cell(0, 0).text = "요구사항 내용"
        table.cell(2, 0).text = "구현방안"
        requirement = {
            "component_id": "API",
            "name": "API Server",
            "layer": "Application Layer",
            "description": "업무 API를 처리",
        }
        arch_doc = {
            "relations": [{"source": "API", "target": "RDBMS", "description": "메타데이터 저장"}],
            "deployment_environment": {"environment": "운영/검증 분리"},
        }

        docx_exporter._fill_arch_requirement_table(table, requirement, arch_doc)

        self.assertEqual(table.cell(0, 1).text, "")
        self.assertEqual(table.cell(1, 0).text, "업무 API를 처리")
        self.assertIn("API Server는 Application Layer에 배치", table.cell(3, 0).text)

    def test_search_router_normalizes_rag_web_both_and_none(self) -> None:
        class FakeQdrant:
            def query_points(self, **kwargs):
                self.kwargs = kwargs
                return type(
                    "Response",
                    (),
                    {
                        "points": [
                            {
                                "id": "rag-1",
                                "score": 0.9,
                                "payload": {"title": "RAG", "text": "rag content"},
                            }
                        ]
                    },
                )()

        def web_provider(query, top_k, filters):
            return [{"id": "web-1", "title": "WEB", "snippet": "web content"}]

        rag = search("query", query_vector=[0.1], rag_client=FakeQdrant())
        web = search("query", search_targets="WEB", web_provider=web_provider)
        both = search(
            "query",
            search_targets="BOTH",
            query_vector=[0.1],
            rag_client=FakeQdrant(),
            web_provider=web_provider,
        )
        none = search("query", search_targets="NONE")

        self.assertEqual(rag["data"]["normalized_results"][0]["source"], "RAG")
        self.assertEqual(rag["data"]["normalized_results"][0]["source_kind"], "RAG")
        self.assertEqual(rag["data"]["search_type"], "RAG")
        self.assertEqual(rag["data"]["results"][0]["metadata"]["title"], "RAG")
        self.assertEqual(web["data"]["normalized_results"][0]["source"], "WEB")
        self.assertEqual(web["data"]["normalized_results"][0]["citation"], "")
        self.assertEqual(web["data"]["search_type"], "WEB")
        self.assertEqual(web["data"]["results"][0]["title"], "WEB")
        self.assertEqual(len(both["data"]["normalized_results"]), 2)
        self.assertEqual(both["data"]["search_type"], "BOTH")
        self.assertEqual(none["data"]["normalized_results"], [])
        self.assertEqual(none["data"]["search_type"], "NONE")

    def test_search_router_builds_embedding_when_query_vector_is_missing(self) -> None:
        captured = {}

        class FakeQdrant:
            def query_points(self, **kwargs):
                captured.update(kwargs)
                return type("Response", (), {"points": []})()

        result = search(
            "공공데이터 표준 용어",
            rag_client=FakeQdrant(),
            embedding_provider=lambda query: [0.1, 0.2, 0.3],
        )

        self.assertTrue(result["success"])
        self.assertEqual(captured["query"], [0.1, 0.2, 0.3])
        self.assertEqual(result["data"]["search_type"], "RAG")

    def test_embedding_writer_upserts_non_functional_requirements_only(self) -> None:
        class FakeQdrant:
            def __init__(self):
                self.created = []
                self.upserts = []

            def collection_exists(self, **kwargs):
                return False

            def create_collection(self, **kwargs):
                self.created.append(kwargs)

            def upsert(self, **kwargs):
                self.upserts.append(kwargs)

        class FakeEmbedder:
            def encode(self, text, normalize_embeddings=True):
                return [0.1, 0.2, 0.3]

        qdrant = FakeQdrant()
        result = write_non_functional_requirements(
            [
                {"req_id": "F-001", "requirement_type": "기능 요구사항", "req_name": "로그인", "detail_text": "로그인한다."},
                {"req_id": "S-001", "requirement_type": "보안 요구사항", "req_name": "계정 잠금", "detail_text": "로그인 실패 시 계정을 잠근다."},
                {"req_id": "D-001", "requirement_type": "데이터 보관", "req_name": "이력 보관", "detail_text": "처리 이력을 보관한다."},
            ],
            project_sn=1,
            source_path="rfp.docx",
            qdrant_client=qdrant,
            embedder=FakeEmbedder(),
            collection="test_collection",
        )

        self.assertTrue(result["success"])
        self.assertEqual(result["data"]["stored_count"], 2)
        self.assertEqual(qdrant.created[0]["collection_name"], "test_collection")
        points = qdrant.upserts[0]["points"]
        payloads = [point.payload for point in points]
        self.assertEqual({payload["requirement_id"] for payload in payloads}, {"S-001", "D-001"})
        self.assertTrue(all(payload["project_sn"] == 1 for payload in payloads))
        self.assertTrue(all("requirement_source_id" in payload for payload in payloads))
        self.assertTrue(all(payload["doc_type"] == "project_non_functional_requirement" for payload in payloads))
        self.assertTrue(all(payload["chunk_type"] == "project_requirement_source" for payload in payloads))

    def test_embedding_writer_defaults_to_alpled_reference_collection(self) -> None:
        class FakeQdrant:
            def __init__(self):
                self.created = []
                self.upserts = []

            def collection_exists(self, **kwargs):
                return False

            def create_collection(self, **kwargs):
                self.created.append(kwargs)

            def upsert(self, **kwargs):
                self.upserts.append(kwargs)

        class FakeEmbedder:
            def encode(self, text, normalize_embeddings=True):
                return [0.1, 0.2, 0.3]

        qdrant = FakeQdrant()
        result = write_non_functional_requirements(
            [{"req_id": "S-001", "requirement_type": "보안", "req_name": "보안", "detail_text": "접근을 제한한다."}],
            project_sn=1,
            source_path="rfp.docx",
            qdrant_client=qdrant,
            embedder=FakeEmbedder(),
        )

        self.assertTrue(result["success"])
        self.assertEqual(result["data"]["collection"], "ALPLED_reference")
        self.assertEqual(qdrant.created[0]["collection_name"], "ALPLED_reference")

    def test_search_request_validates_contract(self) -> None:
        request = SearchRequest(
            project_sn=1,
            docs_cd="SRS",
            agent_name="requirement_generation_agent",
            search_intent="비기능 요구사항 검색",
            query=" requirements ",
            search_targets="BOTH",
            filters={"requirement_type": ["보안 요구사항"]},
            top_k=10,
        )
        self.assertEqual(request.query, "requirements")
        self.assertEqual(request.project_sn, 1)

    def test_search_router_accepts_request_dict_contract(self) -> None:
        def web_provider(query, top_k, filters):
            self.assertEqual(filters["source_type"], ["POLICY"])
            return [
                {
                    "title": "표준",
                    "url": "https://example.test/standard",
                    "snippet": "standard content",
                    "source": "example",
                    "published_at": "2026-01-01",
                }
            ]

        result = search(
            {
                "project_sn": 1,
                "docs_cd": "SRS",
                "agent_name": "requirement_generation_agent",
                "search_intent": "비기능 요구사항 검색",
                "query": "로그인 보안 정책",
                "search_targets": "WEB",
                "filters": {"source_type": ["POLICY"]},
                "top_k": 5,
            },
            web_provider=web_provider,
        )

        self.assertTrue(result["success"])
        self.assertEqual(result["data"]["request"]["agent_name"], "requirement_generation_agent")
        self.assertEqual(result["data"]["normalized_results"][0]["source_kind"], "WEB")
        self.assertEqual(
            result["data"]["normalized_results"][0]["citation"],
            "https://example.test/standard",
        )

    def test_erd_template_diagram_cell_contains_image_without_relationship_text(self) -> None:
        png_bytes = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAFgwJ/lN1mAAAAAABJRU5ErkJggg=="
        )
        with tempfile.TemporaryDirectory() as root:
            root_path = Path(root)
            image_path = root_path / "erd.png"
            image_path.write_bytes(png_bytes)
            output_path = root_path / "erd.docx"

            result = export_docx(
                {
                    "docs_cd": "ERD",
                    "content": {
                        "erd_entity_json": {
                            "erd_id": "ERD-001",
                            "erd_name": "테스트 ERD",
                            "tables": [
                                {
                                    "table_id": "ENT-001",
                                    "logical_name": "너무 긴 사용자 계정 인증 관리 엔티티 이름",
                                    "physical_name": "tbl_user_account",
                                    "description": "사용자 계정 인증 정보를 관리하는 엔티티입니다.",
                                    "columns": [
                                        {
                                            "logical_name": "사용자 계정 일련번호",
                                            "physical_name": "user_account_sn",
                                            "data_type": "BIGINT",
                                            "nullable": False,
                                            "constraints": ["PK"],
                                        }
                                    ],
                                }
                            ],
                            "relationships": [
                                {"from_table": "tbl_user_account", "to_table": "tbl_docs", "description": "relates"}
                            ],
                        }
                    },
                    "image_paths": [str(image_path)],
                },
                str(output_path),
                template_path="templates/erd_template.docx",
            )

            self.assertTrue(result["success"])
            document = Document(str(output_path))
            self.assertNotIn("relates", document.tables[1].cell(1, 0).text)
            self.assertEqual(document.tables[2].cell(0, 7).text.strip(), "USER_ACCOUNT")
            self.assertEqual(document.tables[2].cell(3, 0).text.strip(), "USER_ACCOUNT_SN")

    def test_erd_column_constraint_cell_does_not_fallback_to_description(self) -> None:
        description_only_row = docx_exporter._erd_column_to_row(
            {
                "logical_name": "사용자명",
                "physical_name": "user_nm",
                "data_type": "VARCHAR(100)",
                "nullable": True,
                "constraints": [],
                "description": "사용자 명칭",
            }
        )
        constrained_row = docx_exporter._erd_column_to_row(
            {
                "logical_name": "비밀번호",
                "physical_name": "password_hash",
                "data_type": "VARCHAR(255)",
                "nullable": False,
                "constraints": ["비밀번호는 해시로 저장해야 한다."],
            }
        )

        self.assertEqual(description_only_row[-1], "")
        self.assertEqual(constrained_row[-1], "비밀번호는 해시로 저장해야 한다.")

    def test_docx_export_strips_problematic_png_dpi_metadata(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            root_path = Path(root)
            image_path = root_path / "bad_dpi.png"
            output_path = root_path / "erd.docx"
            from PIL import Image

            Image.new("RGB", (300, 80), "white").save(image_path, dpi=(35.47, 35.47))

            result = export_docx(
                {
                    "docs_cd": "ERD",
                    "content": {
                        "erd_entity_json": {
                            "erd_id": "ERD-001",
                            "erd_name": "테스트 ERD",
                            "tables": [],
                            "relationships": [],
                        }
                    },
                    "image_paths": [str(image_path)],
                },
                str(output_path),
                template_path="templates/erd_template.docx",
            )

            self.assertTrue(result["success"], result.get("error"))
            self.assertTrue(output_path.exists())

    def test_docx_export_sanitizes_decimal_template_widths(self) -> None:
        safe_template = docx_exporter._docx_safe_template_path(Path("templates/erd_template.docx"))
        from zipfile import ZipFile

        with ZipFile(safe_template) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")

        self.assertNotIn('w:w="1396.5"', document_xml)

    def test_mermaid_renderer_uses_high_resolution_options_without_rewriting_dpi(self) -> None:
        captured = {}

        def fake_run(args, **kwargs):
            captured["args"] = args
            image_path = Path(args[args.index("-o") + 1])
            image_path.parent.mkdir(parents=True, exist_ok=True)
            from PIL import Image

            Image.new("RGB", (10, 10), "white").save(image_path)
            return type("Completed", (), {"returncode": 0, "stderr": ""})()

        original_run = mermaid_renderer.subprocess.run
        try:
            mermaid_renderer.subprocess.run = fake_run
            with tempfile.TemporaryDirectory() as root:
                result = mermaid_renderer.render_mermaid(
                    "flowchart TD\nA --> B",
                    output_dir=root,
                )
        finally:
            mermaid_renderer.subprocess.run = original_run

        self.assertTrue(result["success"])
        self.assertEqual(captured["args"][captured["args"].index("-w") + 1], "1400")
        self.assertEqual(captured["args"][captured["args"].index("-H") + 1], "2200")
        self.assertEqual(captured["args"][captured["args"].index("-s") + 1], "2")
        self.assertNotIn("dpi", result["data"]["render_options"])

    def test_llm_client_uses_injected_transport(self) -> None:
        captured = {}

        def transport(url, payload, headers, timeout):
            captured.update(
                {"url": url, "payload": payload, "headers": headers, "timeout": timeout}
            )
            return {"choices": [{"message": {"content": "{}"}}]}

        client = LLMClient(
            base_url="http://llm.test/v1",
            api_key="secret",
            model_name="test-model",
            timeout=10,
            transport=transport,
        )
        result = client.chat([{"role": "user", "content": "hello"}])

        self.assertTrue(result["success"])
        self.assertEqual(captured["url"], "http://llm.test/v1/chat/completions")
        self.assertEqual(captured["payload"]["model"], "test-model")
        self.assertEqual(captured["headers"]["Authorization"], "Bearer secret")
        self.assertEqual(captured["timeout"], 10)

    def test_parallel_llm_calls_preserve_input_order(self) -> None:
        def transport(url, payload, headers, timeout):
            return payload["messages"][0]["content"]

        client = LLMClient(transport=transport)
        result = send_parallel(
            [
                {"messages": [{"role": "user", "content": "first"}]},
                {"messages": [{"role": "user", "content": "second"}]},
            ],
            client=client,
            max_workers=2,
        )

        self.assertTrue(result["success"])
        self.assertEqual(
            [item["data"] for item in result["data"]],
            ["first", "second"],
        )

    def test_json_parser_repairs_code_fence_and_trailing_comma(self) -> None:
        result = repair_json_text('```json\n{"value": 1,}\n```')
        self.assertTrue(result["success"])
        self.assertEqual(result["data"]["value"], {"value": 1})

        parsed = parse_json_response('{"items": []}')
        self.assertTrue(parsed["success"])

    def test_local_download_upload_and_cleanup(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            root_path = Path(root)
            source = root_path / "source.txt"
            source.write_text("hello", encoding="utf-8")

            download = download_file(
                file_path=str(source),
                file_name="copy.txt",
                destination_dir=str(root_path / "input"),
            )
            self.assertTrue(download["success"])

            upload = upload_file(
                download["data"]["local_file_path"],
                storage_path=str(root_path / "output" / "copy.txt"),
            )
            self.assertTrue(upload["success"])

            cleanup = cleanup_paths(
                [download["data"]["local_file_path"], upload["data"]["storage_file_path"]],
                allowed_root=root,
            )
            self.assertTrue(cleanup["success"])

    def test_docx_parser(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            path = Path(root) / "sample.docx"
            document = Document()
            document.add_paragraph("sample text")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "a"
            table.cell(0, 1).text = "b"
            document.save(path)

            result = parse_docx(str(path))
            self.assertTrue(result["success"])
            self.assertEqual(result["data"]["paragraphs"], ["sample text"])
            self.assertEqual(result["data"]["tables"][0][0], ["a", "b"])

            tables = parse_tables(str(path))
            self.assertTrue(tables["success"])
            self.assertEqual(tables["data"]["tables"][0][0], ["a", "b"])

    def test_rfp_rule_parser_supports_injected_parser(self) -> None:
        result = parse_rfp_requirements(
            "sample.docx",
            parser=lambda file_path: [
                {"requirement_id": "SFR-001", "requirement_type": "기능"}
            ],
        )

        self.assertTrue(result["success"])
        self.assertEqual(
            result["data"]["functional_requirements"],
            [{"requirement_id": "SFR-001", "requirement_type": "기능"}],
        )
        self.assertEqual(result["data"]["document_id"], "DOC-001")
        self.assertEqual(result["data"]["document_name"], "sample.docx")

    def test_rfp_rule_parser_extracts_docx_table_requirements(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            path = Path(root) / "rfp.docx"
            document = Document()
            table = document.add_table(rows=5, cols=2)
            table.cell(0, 0).text = "요구사항 고유번호"
            table.cell(0, 1).text = "SFR-001"
            table.cell(1, 0).text = "요구사항명"
            table.cell(1, 1).text = "로그인 기능"
            table.cell(2, 0).text = "요구사항 정의"
            table.cell(2, 1).text = "사용자 인증 기능"
            table.cell(3, 0).text = "요구사항 상세설명"
            table.cell(3, 1).text = "사용자는 아이디와 비밀번호로 로그인할 수 있어야 한다."
            table.cell(4, 0).text = "검증기준"
            table.cell(4, 1).text = "정상 로그인 여부를 확인한다."
            document.save(path)

            result = parse_rfp_requirements(str(path))

            self.assertTrue(result["success"])
            item = result["data"]["functional_requirements"][0]
            self.assertEqual(item["requirement_id"], "SFR-001")
            self.assertEqual(item["requirement_name"], "로그인 기능")
            self.assertEqual(item["requirement_type"], "기능")
            self.assertEqual(item["requirement_definition"], "사용자 인증 기능")
            self.assertIn("로그인", item["requirement_detail"])
            self.assertNotIn("사용자 인증 기능", item["requirement_detail"])
            self.assertEqual(
                item["source_location"],
                {
                    "table_index": 0,
                    "source_type": "detailed_requirement_table",
                },
            )

    def test_rfp_rule_parser_extracts_pdf_text_requirements(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            path = Path(root) / "rfp.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text(
                (72, 72),
                "\n".join(
                    [
                        "SFR-001",
                        "Login feature allows users to sign in with ID and password.",
                        "The system must validate credentials and return the login result.",
                    ]
                ),
            )
            document.save(path)
            document.close()

            result = parse_rfp_requirements(str(path))

            self.assertTrue(result["success"])
            item = result["data"]["functional_requirements"][0]
            self.assertEqual(item["requirement_id"], "SFR-001")
            self.assertIn("Login feature", item["requirement_name"])
            self.assertEqual(item["requirement_definition"], "")
            self.assertIn("credentials", item["requirement_detail"])

    def test_pdf_parser(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            path = Path(root) / "sample.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text((72, 72), "sample pdf")
            document.save(path)
            document.close()

            result = parse_pdf(str(path))
            self.assertTrue(result["success"])
            self.assertIn("sample pdf", result["data"]["text"])

    def test_standard_error_result(self) -> None:
        result = search("test", search_targets="INVALID")  # type: ignore[arg-type]
        self.assertFalse(result["success"])
        self.assertEqual(
            set(result["error"]),
            {"code", "message", "details"},
        )


if __name__ == "__main__":
    unittest.main()
