"""
[PHASE 1-A] 범용 문서 리더 (PDF / DOCX)
==========================================
지원 형식:
  .pdf   → PyMuPDF (fitz) 로 텍스트 직접 추출
  .docx  → python-docx 로 단락·표 구조 보존 추출
  .hwp / .hwpx → ❌ 미지원 (아래 안내 참고)
  .txt / .md   → 직접 읽기

[HWP 미지원 안내]
  LibreOffice를 통한 HWP → DOCX 변환이 실제 환경에서 불안정합니다.
  현재 단계에서는 HWP 파일을 직접 처리하지 않습니다.
  
  대신 아래 방법으로 DOCX로 변환 후 사용하세요:
  
  방법 1 (권장): 한글(HWP)에서 직접 저장
    한글 프로그램 → 파일 → 다른 이름으로 저장 → 파일 형식: MS 워드(*.docx)
  
  방법 2: 플랫폼 변환 (나중에 통합 예정)
    서비스 내 HWP → DOCX 변환 기능 추가 예정

설치 패키지:
  pip install pymupdf python-docx

실행 방법 (단독 테스트):
  python phase1a_doc_reader.py [파일경로]
"""

import os
import re
from pathlib import Path
from typing import Optional


# ─────────────────────────────────────────────
# 텍스트 정제 (노이즈 제거)
# ─────────────────────────────────────────────
def clean_text(text: str) -> str:
    """
    PDF/DOCX에서 추출된 텍스트의 노이즈를 제거합니다.

    제거 대상:
      - 특수문자 반복 라인 (구분선: ────, ====)
      - 3자 이하 단독 줄 (페이지 번호, 머리글/바닥글)
      - 4개 이상 연속 공백 → 단일 공백
      - 3줄 이상 연속 빈 줄 → 2줄로 압축
    """
    # 특수문자 반복 라인 제거
    text = re.sub(r"[─═━\-=\*\.]{4,}", "", text)

    # 3자 이하 단독 줄 제거 (페이지 번호 등)
    lines = []
    for line in text.split("\n"):
        stripped = line.strip()
        if len(stripped) > 3 or stripped == "":
            lines.append(line)
    text = "\n".join(lines)

    # 다중 공백 → 단일 공백
    text = re.sub(r"[ \t]{4,}", " ", text)

    # 3줄 이상 연속 빈 줄 → 2줄
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


# ─────────────────────────────────────────────
# PDF 텍스트 추출
# ─────────────────────────────────────────────
def extract_text_from_pdf(pdf_path: str, max_chars: int = 0) -> str:
    """
    PDF에서 전체 텍스트를 추출합니다.
    스캔된 이미지 PDF는 텍스트 레이어가 없어 빈 결과가 나올 수 있습니다.

    Args:
        pdf_path : PDF 파일 경로
        max_chars: 최대 문자 수 (0이면 전체 추출)
    Returns:
        정제된 텍스트 문자열
    Raises:
        ImportError: pip install pymupdf 필요
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF가 필요합니다: pip install pymupdf")

    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"파일 없음: {pdf_path}")

    doc = fitz.open(pdf_path)
    total_pages = len(doc)
    print(f"   PDF: {total_pages}페이지")

    texts = []
    for page in doc:
        page_text = page.get_text()
        if page_text.strip():
            texts.append(page_text)
    doc.close()

    full_text = "\n\n".join(texts)
    full_text = clean_text(full_text)

    if max_chars > 0 and len(full_text) > max_chars:
        print(f"   ⚠️  {len(full_text):,}자 → {max_chars:,}자로 잘림")
        full_text = full_text[:max_chars]

    return full_text


# ─────────────────────────────────────────────
# DOCX 텍스트 추출
# ─────────────────────────────────────────────
def extract_text_from_docx(docx_path: str, max_chars: int = 0) -> str:
    """
    DOCX에서 단락 + 표 구조를 보존하며 텍스트를 추출합니다.

    Args:
        docx_path: DOCX 파일 경로
        max_chars: 최대 문자 수 (0이면 전체)
    Returns:
        정제된 텍스트 문자열
    Raises:
        ImportError: pip install python-docx 필요
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx가 필요합니다: pip install python-docx")

    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"파일 없음: {docx_path}")

    doc = Document(docx_path)
    parts = []

    # 단락 추출 (헤딩 레벨 감지)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower()
        if "heading" in style:
            level_match = re.search(r"\d+", style)
            level = int(level_match.group()) if level_match else 1
            parts.append(f"{'#' * level} {text}")
        else:
            parts.append(text)

    # 표 추출 (마크다운 형식)
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells]
            # 병합 셀 중복 제거
            unique = []
            for t in row_texts:
                if not unique or unique[-1] != t:
                    unique.append(t)
            parts.append(" | ".join(unique))
        parts.append("")

    full_text = "\n\n".join(parts)
    full_text = clean_text(full_text)

    if max_chars > 0 and len(full_text) > max_chars:
        full_text = full_text[:max_chars]

    return full_text


# ─────────────────────────────────────────────
# 메인 함수: 파일 형식 자동 감지 + 추출
# ─────────────────────────────────────────────
def read_document(file_path: str, max_chars: int = 0) -> str:
    """
    파일 확장자를 자동 감지하여 텍스트를 추출합니다.
    에이전트 코드에서는 항상 이 함수를 호출하면 됩니다.

    지원: .pdf, .docx, .txt, .md
    미지원: .hwp, .hwpx (한글에서 .docx로 저장 후 사용)

    Args:
        file_path: 입력 파일 경로
        max_chars: 최대 추출 문자 수 (0이면 전체)
    Returns:
        정제된 텍스트 문자열
    Raises:
        FileNotFoundError: 파일 없음
        ValueError       : 지원하지 않는 형식
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"파일 없음: {file_path}")

    ext = Path(file_path).suffix.lower()
    file_size = os.path.getsize(file_path)
    print(f"   📄 {Path(file_path).name}  ({file_size:,} bytes, 형식: {ext})")

    if ext == ".pdf":
        return extract_text_from_pdf(file_path, max_chars)

    elif ext == ".docx":
        return extract_text_from_docx(file_path, max_chars)

    elif ext in (".hwp", ".hwpx"):
        raise ValueError(
            "HWP/HWPX 파일은 현재 직접 처리할 수 없습니다.\n\n"
            "  ✅ 해결 방법:\n"
            "     한글 프로그램 → 파일 → 다른 이름으로 저장\n"
            "     → 파일 형식: 'MS 워드(*.docx)' 선택 후 저장\n\n"
            "  변환된 .docx 파일을 다시 전달해주세요."
        )

    elif ext in (".txt", ".md"):
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        text = clean_text(text)
        return text[:max_chars] if max_chars > 0 else text

    else:
        raise ValueError(
            f"지원하지 않는 파일 형식: {ext}\n"
            "지원 형식: .pdf, .docx, .txt, .md\n"
            "HWP 파일은 한글 프로그램에서 .docx로 저장 후 사용하세요."
        )


# ─────────────────────────────────────────────
# 청킹 (RAG용 분할)
# ─────────────────────────────────────────────
def chunk_document(text: str, max_chars: int = 400) -> list[str]:
    """
    추출된 텍스트를 RAG 적재용 청크로 분할합니다.

    전략:
      1. 빈 줄(\\n\\n)로 단락 분리
      2. 단락이 max_chars 초과 시 줄 단위로 재분할
      3. 30자 미만 짧은 청크는 앞 청크에 병합

    Args:
        text     : read_document() 결과
        max_chars: 청크 최대 문자 수
    Returns:
        청크 문자열 리스트
    """
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks = []

    for para in paragraphs:
        if len(para) <= max_chars:
            chunks.append(para)
        else:
            lines = [l.strip() for l in para.split("\n") if l.strip()]
            buffer = ""
            for line in lines:
                if len(buffer) + len(line) + 1 <= max_chars:
                    buffer = (buffer + " " + line).strip()
                else:
                    if buffer:
                        chunks.append(buffer)
                    buffer = line
            if buffer:
                chunks.append(buffer)

    # 30자 미만 청크 앞에 병합
    merged = []
    for chunk in chunks:
        if len(chunk) < 30 and merged:
            merged[-1] += " " + chunk
        else:
            merged.append(chunk)

    return merged


# ─────────────────────────────────────────────
# 단독 실행 테스트
# ─────────────────────────────────────────────
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("사용법: python phase1a_doc_reader.py [파일경로]")
        print("예시:   python phase1a_doc_reader.py data/RFP_원본.pdf")
        sys.exit(0)

    target = sys.argv[1]
    print(f"\n[테스트] {target} 읽기 시작...")

    text = read_document(target, max_chars=3000)
    print(f"\n✅ 추출 완료! ({len(text):,}자)\n")
    print("--- 앞 500자 미리보기 ---")
    print(text[:500])
    print("...")

    chunks = chunk_document(text)
    print(f"\n📦 청킹 결과: {len(chunks)}개 청크")
    for i, c in enumerate(chunks[:3], 1):
        print(f"  [{i}] {c[:80]}...")
