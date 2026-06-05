"""
통합시험 시나리오 생성 에이전트 테스트 코드 (RunPod 환경)

사용법:
    # UI 설계서 없이 실행 (요구사항 정의서만)
    python TS_agent.py \
        --model exaone \
        --input 요구사항_정의서.json

    # UI 설계서 포함 실행 (파트 2 JSON 파일들을 --ui 로 전달)
    python TS_agent.py \
        --model exaone \
        --input 요구사항_정의서.json \
        --ui UI_설계서.json

실행 결과:
    - 요구사항_정의서_output.json  : 통합시험 시나리오 JSON
    - 요구사항_정의서_output.docx  : CBD 포맷 Word 문서 (자동 생성)

지원 모델:
    - exaone   : LGAI-EXAONE/EXAONE-3.5-7.8B-Instruct (권장)
    - exaone2b : LGAI-EXAONE/EXAONE-3.5-2.4B-Instruct (로컬 테스트용)
    - qwen     : Qwen/Qwen2.5-7B-Instruct
    - qwen3b   : Qwen/Qwen2.5-3B-Instruct (로컬 테스트용)

RunPod 세팅 가이드:
    1. RunPod 접속 후 GPU 선택
       - RTX 3090 / RTX 4090 (24GB) 이상 권장
    2. Template: "Ubuntu 22.04" 선택
    3. Pod 생성 후 Web Terminal 또는 SSH 접속
    4. Ollama 설치:
       curl -fsSL https://ollama.com/install.sh | sh
    5. Ollama 서버 실행:
       ollama serve &
    6. 모델 다운로드:
       ollama pull exaone3.5:7.8b
       ollama pull qwen2.5:7b
    7. 의존성 설치:
       pip install ollama python-docx
    8. 파일 업로드 후 실행:
       python TS_agent.py --model exaone --input 요구사항_정의서.json
"""

import argparse
import json
import sys
import time
from pathlib import Path

# ─────────────────────────────────────────────
# 프롬프트 import (같은 디렉토리에 있어야 함)
# ─────────────────────────────────────────────
try:
    from TS_prompt import SYSTEM_PROMPT, build_prompt
except ImportError:
    try:
        from TS.TS_prompt import SYSTEM_PROMPT, build_prompt
    except ImportError:
        print("[ERROR] TS_prompt.py 파일이 같은 디렉토리에 있어야 합니다.")
        sys.exit(1)


# ─────────────────────────────────────────────
# 모델별 클라이언트 초기화
# ─────────────────────────────────────────────

def get_client(model_type: str, model_name: str | None = None):
    """모델 타입에 따라 클라이언트를 반환합니다."""

    try:
        import ollama
    except ImportError:
        print("[ERROR] 의존성 설치 필요:")
        print("        pip install ollama")
        sys.exit(1)

    if model_type == "exaone":
        actual_model = model_name or "exaone3.5:7.8b"
    elif model_type == "exaone2b":
        actual_model = model_name or "exaone3.5:2.4b"
    elif model_type == "qwen":
        actual_model = model_name or "qwen2.5:7b"
    elif model_type == "qwen3b":
        actual_model = model_name or "qwen2.5:3b"
    else:
        print(f"[ERROR] 지원하지 않는 모델 타입: {model_type}")
        print("        지원 모델: exaone, exaone2b, qwen, qwen3b")
        sys.exit(1)

    print(f"[INFO] 모델 로딩 중: {actual_model}")
    return {"type": "ollama", "model": actual_model}


# ─────────────────────────────────────────────
# 모델 추론
# ─────────────────────────────────────────────

def run_inference(client: dict, messages: list, system_prompt: str) -> str:
    """추론을 실행하고 응답 텍스트를 반환합니다."""
    import ollama

    model = client["model"]

    # system + few-shot + 실제 input을 메시지 구성
    full_messages = [{"role": "system", "content": system_prompt}] + messages

    print(f"[INFO] 입력 메시지 수: {len(full_messages)}")

    response = ollama.chat(
        model=model,
        messages=full_messages,
        format="json",
        options={
            "temperature": 0,        # greedy decoding: 재현성 확보
            "num_predict": 16384,    # max_new_tokens에 대응
        }
    )

    output_text = response["message"]["content"]
    print(f"[INFO] 출력 완료")

    return output_text


# ─────────────────────────────────────────────
# JSON 파싱 및 검증
# ─────────────────────────────────────────────

def parse_and_validate(raw_output: str) -> tuple[dict | None, str]:
    """
    모델 출력에서 JSON을 파싱하고 스키마 기본 검증을 수행합니다.

    Returns:
        (parsed_dict, error_message)
        성공 시 error_message는 빈 문자열
    """
    # 마크다운 코드블록 제거 (모델이 출력 규칙을 무시하는 경우 대비)
    text = raw_output.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        text = "\n".join(lines[1:])
    if text.endswith("```"):
        text = text[:-3].strip()

    # JSON 파싱
    try:
        data = json.loads(text)
    except json.JSONDecodeError as e:
        return None, f"JSON 파싱 실패: {e}"

    # 필수 최상위 키 확인
    if "scenarios" not in data:
        return None, "필수 키 누락: 'scenarios'"
    if "cases" not in data:
        return None, "필수 키 누락: 'cases'"

    # scenarios 구조 확인
    for i, scenario in enumerate(data["scenarios"]):
        required = ["scenario_id", "scenario_name", "scenario_description", "test_cases"]
        for key in required:
            if key not in scenario:
                return None, f"scenarios[{i}] 필수 키 누락: '{key}'"
        for j, tc in enumerate(scenario["test_cases"]):
            required_tc = ["test_case_id", "test_case_description", "test_procedure", "scenario_detail"]
            for key in required_tc:
                if key not in tc:
                    return None, f"scenarios[{i}].test_cases[{j}] 필수 키 누락: '{key}'"

    # cases 구조 확인
    for i, case in enumerate(data["cases"]):
        required = ["round", "scenario_id", "scenario_name", "test_case_id",
                    "sequence", "process_content", "test_item",
                    "input_data", "expected_result", "screen_id"]
        for key in required:
            if key not in case:
                return None, f"cases[{i}] 필수 키 누락: '{key}'"

        # test_result는 설계단계에서 null이어야 함
        # sLLM이 hallucination으로 임의 값을 채우는 경우를 감지
        if case.get("test_result") is not None:
            return None, f"cases[{i}].test_result는 설계단계에서 null이어야 합니다. (hallucination 감지)"

        # input_data 타입 검증
        if not isinstance(case.get("input_data"), str):
            return None, f"cases[{i}].input_data는 문자열이어야 합니다. (현재: {type(case.get('input_data')).__name__})"

    # test_procedure ↔ cases 정합성 확인 (경고만 출력, 실패 처리하지 않음)
    for scenario in data["scenarios"]:
        for tc in scenario["test_cases"]:
            tc_id = tc["test_case_id"]
            procedure_count = len(tc["test_procedure"])
            case_count = sum(1 for c in data["cases"] if c["test_case_id"] == tc_id)
            if procedure_count != case_count:
                print(f"[WARN] {tc_id}: test_procedure 항목 수({procedure_count})와 "
                      f"cases 행 수({case_count})가 일치하지 않습니다.")

    return data, ""


# ─────────────────────────────────────────────
# cases 누락 후처리 보완
# ─────────────────────────────────────────────

def fill_missing_cases(data: dict) -> dict:
    """
    test_procedure 항목 수 대비 cases 행이 부족한 경우 누락된 행을 자동으로 보완합니다.

    보완 전략:
    - 해당 test_case_id의 기존 cases 중 sequence가 가장 높은 행을 템플릿으로 사용
    - 누락된 sequence에 대해 process_content를 test_procedure 항목으로 채워 행을 추가
    - test_item, input_data, expected_result는 "(자동 보완 필요)" 로 표시하여
      사람이 나중에 검토할 수 있도록 함
    """
    filled_count = 0

    for scenario in data["scenarios"]:
        scenario_id = scenario["scenario_id"]
        scenario_name = scenario["scenario_name"]

        for tc in scenario["test_cases"]:
            tc_id = tc["test_case_id"]
            procedures = tc["test_procedure"]
            procedure_count = len(procedures)

            # 현재 생성된 cases에서 이 tc_id의 행들을 sequence 기준으로 수집
            existing = {
                c["sequence"]: c
                for c in data["cases"]
                if c["test_case_id"] == tc_id
            }

            if len(existing) >= procedure_count:
                continue  # 이미 충분하면 스킵

            # 템플릿으로 쓸 기존 case (sequence 가장 높은 것)
            template = existing[max(existing.keys())] if existing else None

            for seq in range(1, procedure_count + 1):
                if seq in existing:
                    continue  # 이미 있는 sequence는 스킵

                # 누락된 행 생성
                new_case = {
                    "round": template["round"] if template else 1,
                    "scenario_id": scenario_id,
                    "scenario_name": scenario_name,
                    "test_case_id": tc_id,
                    "sequence": seq,
                    "process_content": procedures[seq - 1],
                    "test_item": "(자동 보완 필요)",
                    "precondition": None,
                    "input_data": "(자동 보완 필요)",
                    "expected_result": "(자동 보완 필요)",
                    "screen_id": template["screen_id"] if template else "",
                    "test_result": None,
                    "note": "자동 보완된 행입니다. 내용을 검토하고 수정하세요."
                }
                data["cases"].append(new_case)
                filled_count += 1
                print(f"[FIX] {tc_id} sequence {seq} 자동 보완")

    if filled_count > 0:
        # sequence 순서대로 정렬 (scenario_id → test_case_id → sequence)
        data["cases"].sort(key=lambda c: (c["scenario_id"], c["test_case_id"], c["sequence"]))
        print(f"[INFO] 총 {filled_count}개 행 자동 보완 완료")
    else:
        print(f"[INFO] cases 누락 없음. 자동 보완 불필요.")

    return data



# ─────────────────────────────────────────────
# docx 변환
# ─────────────────────────────────────────────

def convert_to_docx(data: dict, output_path: str):
    """
    통합시험 시나리오 JSON을 CBD 포맷 Word 문서(.docx)로 변환합니다.

    각 시나리오마다 시험 시나리오 표 + 시험 케이스 표를 생성합니다.
    자동 보완된 행(note에 '자동 보완' 포함)은 빨간색으로 표시됩니다.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy
    except ImportError:
        print("[ERROR] python-docx 설치 필요: pip install python-docx")
        return

    def set_cell_bg(cell, hex_color):
        """셀 배경색 설정"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def set_col_widths(table, widths_cm):
        """컬럼 너비 설정"""
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(widths_cm):
                    cell.width = Cm(widths_cm[i])

    def add_header_row(table, headers, widths_cm):
        """헤더 행 추가"""
        row = table.rows[0]
        for i, (header, width) in enumerate(zip(headers, widths_cm)):
            cell = row.cells[i]
            cell.width = Cm(width)
            set_cell_bg(cell, "4472C4")
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = "맑은 고딕"
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def add_data_cell(cell, text, width_cm, center=False, is_auto=False):
        """데이터 셀 채우기"""
        val = "" if text is None else str(text)
        cell.width = Cm(width_cm)
        para = cell.paragraphs[0]
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(val)
        run.font.size = Pt(8.5)
        run.font.name = "맑은 고딕"
        if is_auto:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc = Document()

    # 페이지 설정: A4 가로
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # 제목
    title = doc.add_heading("통합시험 시나리오", level=1)
    title.runs[0].font.name = "맑은 고딕"

    for scenario in data["scenarios"]:
        # 시나리오 소제목
        heading = doc.add_heading(
            f"{scenario['scenario_id']} - {scenario['scenario_name']}", level=2
        )
        heading.runs[0].font.name = "맑은 고딕"

        # ── 시험 시나리오 표 ──
        p = doc.add_paragraph("■ 시험 시나리오")
        p.runs[0].bold = True
        p.runs[0].font.name = "맑은 고딕"
        p.runs[0].font.size = Pt(10)

        sc_headers = ["시험시나리오 ID", "시험시나리오명", "시험시나리오 설명",
                      "시험케이스 ID", "시험케이스 설명", "시험 절차", "시나리오 설명", "비고"]
        sc_widths  = [2.8, 3.2, 4.5, 2.5, 3.5, 4.5, 4.0, 1.5]
        num_tc = len(scenario["test_cases"])

        sc_table = doc.add_table(rows=1 + num_tc, cols=len(sc_headers))
        sc_table.style = "Table Grid"
        add_header_row(sc_table, sc_headers, sc_widths)

        for ti, tc in enumerate(scenario["test_cases"]):
            row = sc_table.rows[1 + ti]
            procedure_text = "\n".join(
                f"{i+1}. {p}" for i, p in enumerate(tc["test_procedure"])
            )
            values = [
                scenario["scenario_id"] if ti == 0 else "",
                scenario["scenario_name"] if ti == 0 else "",
                scenario["scenario_description"] if ti == 0 else "",
                tc["test_case_id"],
                tc["test_case_description"],
                procedure_text,
                tc.get("scenario_detail", ""),
                tc.get("note", "") or "",
            ]
            centers = [True, False, False, True, False, False, False, False]
            for ci, (val, w, center) in enumerate(zip(values, sc_widths, centers)):
                add_data_cell(row.cells[ci], val, w, center=center)

        doc.add_paragraph()

        # ── 시험 케이스 표 ──
        p2 = doc.add_paragraph("■ 시험 케이스")
        p2.runs[0].bold = True
        p2.runs[0].font.name = "맑은 고딕"
        p2.runs[0].font.size = Pt(10)

        case_headers = ["차수", "시험케이스 ID", "순번", "업무처리 내용", "시험항목",
                        "사전조건", "입력자료", "예상결과", "화면 ID", "시험결과", "비고"]
        case_widths  = [1.0, 2.8, 0.9, 3.5, 3.0, 2.5, 3.0, 3.5, 2.0, 2.0, 2.3]

        scenario_cases = sorted(
            [c for c in data["cases"] if c["scenario_id"] == scenario["scenario_id"]],
            key=lambda c: (c["test_case_id"], c["sequence"])
        )

        ct = doc.add_table(rows=1 + len(scenario_cases), cols=len(case_headers))
        ct.style = "Table Grid"
        add_header_row(ct, case_headers, case_widths)

        for ri, case in enumerate(scenario_cases):
            row = ct.rows[1 + ri]
            is_auto = bool(case.get("note") and "자동 보완" in case["note"])
            values = [
                str(case["round"]),
                case["test_case_id"],
                str(case["sequence"]),
                case["process_content"],
                case["test_item"],
                case.get("precondition") or "",
                case["input_data"],
                case["expected_result"],
                case.get("screen_id") or "",
                "",  # 시험결과 미기술
                case.get("note") or "",
            ]
            centers = [True, True, True, False, False, False, False, False, True, False, False]
            for ci, (val, w, center) in enumerate(zip(values, case_widths, centers)):
                add_data_cell(row.cells[ci], val, w, center=center, is_auto=is_auto)

        doc.add_paragraph()

    doc.save(output_path)
    print(f"[INFO] docx 저장 완료: {output_path}")


# ─────────────────────────────────────────────
# 결과 출력 요약
# ─────────────────────────────────────────────

def print_summary(data: dict):
    """파싱된 결과를 요약 출력합니다."""
    print("\n" + "="*50)
    print("생성 결과 요약")
    print("="*50)

    for scenario in data["scenarios"]:
        print(f"\n[시나리오] {scenario['scenario_id']} - {scenario['scenario_name']}")
        print(f"  설명: {scenario['scenario_description']}")
        print(f"  시험케이스 수: {len(scenario['test_cases'])}개")
        for tc in scenario["test_cases"]:
            print(f"    - {tc['test_case_id']}: {tc['test_case_description']}")
            print(f"      절차 수: {len(tc['test_procedure'])}개")

    print(f"\n[cases] 총 {len(data['cases'])}행")
    from collections import Counter
    counter = Counter(c["test_case_id"] for c in data["cases"])
    for tc_id, count in counter.items():
        print(f"  {tc_id}: {count}행")


# ─────────────────────────────────────────────
# 메인
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="통합시험 시나리오 생성 에이전트 테스트 (RunPod)")
    parser.add_argument("--model", required=True,
                        choices=["exaone", "exaone2b", "qwen", "qwen3b"],
                        help="사용할 모델 타입")
    parser.add_argument("--model-name", default=None,
                        help="모델 이름 직접 지정 (기본값: 모델 타입별 기본값 사용)")
    parser.add_argument("--input", required=True,
                        help="요구사항 정의서 JSON 파일 경로")
    parser.add_argument("--ui", nargs="*", default=None,
                        help="UI 설계서 파트 2 JSON 파일 경로 목록 (여러 개 가능). "
                             "예: --ui UI-001.json UI-002.json")
    parser.add_argument("--output", default=None,
                        help="생성 결과 저장 경로 (기본값: input 파일명_output.json)")
    args = parser.parse_args()

    # ── 1. input 파일 로드
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"[ERROR] input 파일을 찾을 수 없습니다: {input_path}")
        sys.exit(1)

    with open(input_path, "r", encoding="utf-8") as f:
        input_data = json.load(f)

    requirements = input_data.get("requirements", [])
    print(f"[INFO] input 파일 로드 완료: {input_path} (요구사항 {len(requirements)}개)")

    # ── 2. UI 설계서 파일 로드 (있는 경우)
    ui_screens_raw: list[str] | None = None
    if args.ui:
        ui_screens_raw = []
        for ui_path_str in args.ui:
            ui_path = Path(ui_path_str)
            if not ui_path.exists():
                print(f"[WARN] UI 설계서 파일을 찾을 수 없습니다: {ui_path} (건너뜀)")
                continue
            with open(ui_path, "r", encoding="utf-8") as f:
                ui_screens_raw.append(f.read())
        print(f"[INFO] UI 설계서 파일 로드 완료: {len(ui_screens_raw)}개")
    else:
        print("[INFO] UI 설계서 없이 실행합니다.")

    # ── 3. 클라이언트 초기화
    client = get_client(args.model, args.model_name)

    # ── 4. 요구사항 1개씩 순차 처리
    all_scenarios = []
    all_cases = []
    total_start = time.time()

    for i, req in enumerate(requirements):
        req_id = req.get("requirement_id", f"REQ-{i+1}")
        print(f"\n[INFO] [{i+1}/{len(requirements)}] {req_id} 처리 중...")

        # 요구사항 1개짜리 JSON 구성
        single_req_json = json.dumps({"requirements": [req]}, ensure_ascii=False, indent=2)

        # 프롬프트 구성 (UI 설계서 포함 여부에 따라 자동 분기)
        messages = build_prompt(single_req_json, ui_screens_raw)

        # 추론 실행
        start = time.time()
        raw_output = run_inference(client, messages, SYSTEM_PROMPT)
        elapsed = time.time() - start
        print(f"[INFO] 추론 완료 ({elapsed:.1f}초)")

        # JSON 파싱 및 검증
        parsed, error = parse_and_validate(raw_output)

        if error:
            print(f"[FAIL] {req_id} 검증 실패: {error}")
            raw_path = input_path.stem + f"_{req_id}_raw_output.txt"
            with open(raw_path, "w", encoding="utf-8") as f:
                f.write(raw_output)
            print(f"[INFO] raw output 저장됨: {raw_path}")
            continue  # 실패한 요구사항은 건너뛰고 다음 진행

        print(f"[PASS] {req_id} 검증 통과")
        # cases 누락 후처리 보완
        parsed = fill_missing_cases(parsed)
        all_scenarios.extend(parsed.get("scenarios", []))
        all_cases.extend(parsed.get("cases", []))

    total_elapsed = time.time() - total_start
    print(f"\n[INFO] 전체 처리 완료 ({total_elapsed:.1f}초)")
    print(f"[INFO] 성공: 시나리오 {len(all_scenarios)}개, 케이스 {len(all_cases)}개")

    if not all_scenarios:
        print("[ERROR] 생성된 시나리오가 없습니다.")
        sys.exit(1)

    # ── 5. 전체 결과 합치기
    final_output = {
        "scenarios": all_scenarios,
        "cases": all_cases
    }

    # ── 6. 결과 요약 출력
    print_summary(final_output)

    # ── 7. output 저장 (JSON)
    output_path = args.output or (input_path.stem + "_output.json")
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(final_output, f, ensure_ascii=False, indent=2)
    print(f"\n[INFO] JSON 저장 완료: {output_path}")

    # ── 8. docx 변환
    docx_path = str(Path(output_path).with_suffix(".docx"))
    convert_to_docx(final_output, docx_path)


if __name__ == "__main__":
    main()


# ─────────────────────────────────────────────
# cases 누락 후처리 보완 (위치: parse_and_validate 뒤, print_summary 앞)
# ─────────────────────────────────────────────
