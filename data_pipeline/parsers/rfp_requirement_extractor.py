import re
import os
from typing import List, Dict, Any
import docx
from docx.table import Table
from docx.text.paragraph import Paragraph  
from docx.document import Document as DocType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


# =========================
# LEAF REQ DETECTION (변경 핵심)
# =========================
PATTERN_STR = r"\b((?:sfr|sir|cor|cmr|fqr|sec|per|ast|gcl|isr|dar|wtr|uor|prm|ops|mng|csr|ecr|ter|ser|inr|qur|pmr|psr)-?\d{3})"

REQ_ID_PATTERN = re.compile(PATTERN_STR, re.IGNORECASE)


def _extract_req_id(text: str):
    if not text:
        return None
    m = REQ_ID_PATTERN.search(text)
    return m.group(0).upper() if m else None


def _clean(t: str):
    return (t or "").replace("\n", " ").replace("\r", " ").strip()


# =========================
# BULLET PARSER (그대로 유지)
# =========================
def _is_bullet(line: str):
    return bool(re.match(r"^[■▪●○·\-*0-9]+\)", line.strip()))


def _merge_bullets(text: str):
    lines = text.split("\n")
    out, cur = [], ""

    for l in lines:
        l = l.strip()
        if not l:
            continue

        if _is_bullet(l):
            if cur:
                out.append(cur)
            cur = l
        else:
            cur += " " + l if cur else l

    if cur:
        out.append(cur)

    return out


# =========================
# DOCX ITERATOR (그대로 유지)
# =========================
def iter_blocks(doc: DocType):
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


# =========================
# MAIN EXTRACTOR (LOSS 0 → LEAF VERSION)
# =========================
def extract_requirements_from_rfp_docx(file_path: str) -> List[Dict[str, Any]]:

    doc = docx.Document(file_path)

    req_buffer = {}
    results = []

    current_req_id = None

    def flush():
        nonlocal current_req_id, req_buffer
        if current_req_id and req_buffer:
            results.append(req_buffer.copy())
        req_buffer = {}

    for block in iter_blocks(doc):

        # =========================
        # TABLE
        # =========================
        if isinstance(block, Table):
            table_text = "\n".join(
                [" | ".join([_clean(c.text) for c in row.cells]) for row in block.rows]
            )

            req_id = _extract_req_id(table_text)

            # leaf req start
            if req_id:
                if req_id != current_req_id:
                    flush()
                    current_req_id = req_id
                    req_buffer = {
                        "req_id": req_id,
                        "requirement_type": (
                            "기능" if req_id.startswith(("SFR", "CSR")) else "비기능"
                        ),
                        "definition": "",
                        "sub_details": [],
                        "deliverables": [],
                        "constraints": [],
                        "raw_text": table_text
                    }

            if current_req_id:
                req_buffer["raw_text"] += "\n" + table_text
                req_buffer["sub_details"].extend(_merge_bullets(table_text))

            continue

        # =========================
        # PARAGRAPH
        # =========================
        text = _clean(block.text)

        if len(text) < 3:
            continue

        req_id = _extract_req_id(text)

        # NEW leaf requirement
        if req_id:
            if req_id != current_req_id:
                flush()
                current_req_id = req_id
                req_buffer = {
                    "req_id": req_id,
                    "requirement_type": (
                        "기능" if req_id.startswith(("SFR", "CSR")) else "비기능"
                    ),
                    "definition": "",
                    "sub_details": [],
                    "deliverables": [],
                    "constraints": [],
                    "raw_text": text
                }

        # attach
        if current_req_id:
            if "정의" in text or "설명" in text:
                req_buffer["definition"] += "\n" + text
            elif "산출" in text:
                req_buffer["deliverables"].append(text)
            elif "제약" in text:
                req_buffer["constraints"].append(text)
            else:
                req_buffer["sub_details"].append(text)

    flush()

    # =========================
    # POST CLEANING
    # =========================
    cleaned = []
    seen = set()

    for r in results:
        rid = r.get("req_id")

        if not rid or rid in seen:
            continue
        seen.add(rid)

        r["sub_details"] = [x for x in r["sub_details"] if len(x) > 3]
        cleaned.append(r)

    return cleaned


# =========================
# CONTEXT EXTRACTOR (유지)
# =========================
def extract_context_text_from_rfp_docx(file_path: str):
    doc = docx.Document(file_path)
    return [p.text.strip() for p in doc.paragraphs if len(p.text.strip()) > 20]