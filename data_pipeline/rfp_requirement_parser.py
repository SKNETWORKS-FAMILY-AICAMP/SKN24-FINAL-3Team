import os
import re
import json
import docx
from pydantic import BaseModel
from typing import List, Optional

# ======================================================
# Canonical Schema (고정)
# ======================================================

class CanonicalRequirement(BaseModel):
    requirement_id: str
    requirement_name: str
    requirement_type: str
    description: str
    source: List[str]
    constraints: List[str] = []
    priority: str = "미지정"
    validation_criteria: List[str] = []
    note: Optional[str] = None


class RequirementDocument(BaseModel):
    requirements: List[CanonicalRequirement]


# ======================================================
# Config
# ======================================================

ID_PATTERN = re.compile(r"^[A-Za-z]{2,5}[\-\–\—]\d{2,4}$")

DEFAULT_PREFIX_MAP = {
"SFR": "기능", "NFR": "비기능",
        "PER": "성능", "SER": "보안",
        "QUR": "품질", "TER": "테스트",
        "DAR": "데이터", "CSR": "컨설팅",
        "PSR": "프로젝트 지원", "PMR": "프로젝트 관리",
        "COR": "제약사항", "INT": "인터페이스", 
        "SIR": "시스템 장비 구성", "MPR": "유지보수"
}

FIELD_ALIASES = {
    "name": [
        "요구사항명", "요구사항 명", "요구사항 명칭",
        "기능명", "기능 명", "항목명", "업무명"
    ],
    "type": [
        "요구사항분류", "요구사항 구분",
        "구분", "분류", "유형"
    ],
    "constraint": [
        "제약사항", "특이사항", "조건", "비고"
    ],
    "validation": [
        "검증기준", "시험기준", "평가기준", "검수기준"
    ],
    "priority": [
        "우선순위", "중요도"
    ]
}

BLACKLIST = {
    "요구사항명",
    "요구사항 명",
    "요구사항 명칭",
    "요구사항 고유번호",
    "요구사항 상세설명",
    "상세설명",
    "상세 설명",
    "세부내용",
    "세부 내용",
    "정의",
    "산출정보",
    "관련 요구사항",
    "관련요구사항",
    "검토기준",
    "비고",
    "구분",
    "분류",
    "유형",
    "사업수행계획서",
    "설계서",
    "결과보고서"
}


# ======================================================
# Utility
# ======================================================

def load_prefix_map():
    if os.path.exists("requirement_type_map.json"):
        with open("requirement_type_map.json", "r", encoding="utf-8") as f:
            return json.load(f)
    return DEFAULT_PREFIX_MAP


PREFIX_MAP = load_prefix_map()


def normalize_text(text):
    text = text.replace("\n", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def detect_field(text):

    text = normalize_text(text)

    for field, aliases in FIELD_ALIASES.items():
        for alias in aliases:
            if text == alias or alias in text:
                return field

    return None


def infer_requirement_name(parts):

    for p in parts:

        p = normalize_text(p)

        if len(p) < 4:
            continue

        if p.isdigit():
            continue

        if p in BLACKLIST:
            continue

        if "요구사항" in p and len(p) < 20:
            continue

        return p[:100]

    return "미분류"


def infer_requirement_type(req_id, name, desc):

    name = normalize_text(name)

    # 1순위 : 이름

    if "보안" in name:
        return "보안"

    if "성능" in name:
        return "성능"

    if "품질" in name:
        return "품질"

    if "인터페이스" in name:
        return "사용자 인터페이스"

    if "데이터" in name:
        return "데이터"

    if "프로젝트 지원" in name:
        return "프로젝트 지원"

    if "지원" in name:
        return "프로젝트 지원"

    if "장비구성" in name:
        return "시스템 장비구성"

    # 2순위 : Prefix

    prefix = req_id.split("-")[0].upper()

    if prefix in PREFIX_MAP:
        return PREFIX_MAP[prefix]

    # 3순위 : 내용

    text = (name + " " + desc[:500]).lower()

    if any(k in text for k in [
        "보안", "암호화", "접근제어",
        "접근 통제", "개인정보"
    ]):
        return "보안"

    if any(k in text for k in [
        "처리량", "응답속도",
        "동시접속", "throughput"
    ]):
        return "성능"

    if any(k in text for k in [
        "인프라", "서버", "아키텍처"
    ]):
        return "시스템 장비구성"

    return "기능"


# ======================================================
# Parser
# ======================================================

def extract_from_file(file_path):

    doc = docx.Document(file_path)

    requirements_map = {}

    for table in doc.tables:

        last_id = None

        for row in table.rows:

            seen = set()
            cells = []

            for cell in row.cells:

                txt = normalize_text(cell.text)

                if not txt:
                    continue

                if txt in seen:
                    continue

                seen.add(txt)
                cells.append(txt)

            if not cells:
                continue

            id_match = None

            for c in cells:
                clean_c = c.strip()
                if ID_PATTERN.match(clean_c):
                    # 특수 대시를 표준 하이픈('-')으로 치환하고 대문자 통일
                    id_match = re.sub(r"[\-\–\—]", "-", clean_c).upper()
                    break

            # 신규 Requirement

            if id_match:

                last_id = id_match

                if last_id not in requirements_map:

                    requirements_map[last_id] = {
                        "id": last_id,
                        "name": None,
                        "type": None,
                        "priority": None,
                        "constraints": [],
                        "validation_criteria": [],
                        "desc_parts": [],
                        "source": os.path.basename(file_path)
                    }

                remain = [c for c in cells if c != id_match]

                requirements_map[last_id]["desc_parts"].extend(remain)

                continue

            if not last_id:
                continue

            current = requirements_map[last_id]

            if len(cells) >= 2:

                header = cells[0]
                value = " ".join(cells[1:]).strip()

                if "산출정보" in header:
                    continue

                field = detect_field(header)

                if field == "name":
                    current["name"] = value
                    continue

                elif field == "type":
                    current["type"] = value
                    continue

                elif field == "constraint":
                    current["constraints"].append(value)
                    continue

                elif field == "validation":
                    current["validation_criteria"].append(value)
                    continue

                elif field == "priority":
                    current["priority"] = value
                    continue

            clean_parts = []

            for c in cells:

                if c in BLACKLIST:
                    continue

                clean_parts.append(c)

            current["desc_parts"].extend(clean_parts)

    # ==================================================
    # Build Canonical
    # ==================================================

    final_reqs = []

    for data in requirements_map.values():

        req_id = data["id"]

        # 🎯 [추가] ID가 -000 또는 -00으로 끝나는 표 헤더/소개용 행은 통째로 제외
        if req_id.endswith("-000") or req_id.endswith("-00"):
            continue

        unique_parts = []
        seen = set()

        for p in data["desc_parts"]:

            p = normalize_text(p)

            if not p:
                continue

            if p in BLACKLIST:
                continue

            if len(p) <= 2:
                continue

            if p in seen:
                continue

            seen.add(p)
            unique_parts.append(p)

        if not unique_parts:
            continue

        name = data["name"]

        if not name:
            name = infer_requirement_name(unique_parts)

        description = "\n".join(unique_parts)

        description = re.sub(
            r"(요구사항\s*고유번호|요구사항\s*상세설명|세부내용|정의)",
            "",
            description
        )

        description = re.sub(r"\n{2,}", "\n", description).strip()

        if len(description) <= 20:
            continue

        req_type = data["type"]

        if not req_type:
            req_type = infer_requirement_type(
                req_id,
                name,
                description
            )

        final_reqs.append(
            CanonicalRequirement(
                requirement_id=req_id,
                requirement_name=name[:100],
                requirement_type=req_type,
                description=description,
                source=[data["source"]],
                constraints=list(dict.fromkeys(data["constraints"])),
                priority=data["priority"] or "미지정",
                validation_criteria=(
                    data["validation_criteria"]
                    if data["validation_criteria"]
                    else ["검토 필요"]
                )
            )
        )

    return final_reqs


# ======================================================
# Main
# ======================================================

if __name__ == "__main__":

    src_folder_path = r"C:\skn24\수업자료\08_large_language_model\00.final_project\data_pipeline\data\RFP"
    target_folder_path = r"C:\ALPRED_BACKUP\data\requirement_sources"

    os.makedirs(target_folder_path, exist_ok=True)

    for filename in os.listdir(src_folder_path):

        if not filename.endswith(".docx"):
            continue

        try:

            reqs = extract_from_file(
                os.path.join(src_folder_path, filename)
            )

            result = RequirementDocument(
                requirements=reqs
            )

            output_path = os.path.join(
                target_folder_path,
                f"{os.path.splitext(filename)[0]}_전처리이후.json"
            )

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(
                    result.model_dump_json(indent=4)
                )

            print(f"완료 : {filename} ({len(reqs)}건)")

        except Exception as e:
            print(f"실패 : {filename} / {e}")
