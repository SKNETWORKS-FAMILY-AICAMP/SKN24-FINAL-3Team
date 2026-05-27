import os
import re
import uuid
from pathlib import Path
from typing import Any, Dict, Iterable, List

import docx
import fitz
from dotenv import load_dotenv
from qdrant_client.models import PointStruct
from tqdm import tqdm

from rag.qdrant_config import (
    REQUIREMENT_REFERENCE_COLLECTION,
    ensure_named_collection,
    get_client,
    get_embedder,
)

load_dotenv()

COLLECTION_NAME = REQUIREMENT_REFERENCE_COLLECTION
DATA_ROOT = Path(os.getenv("REQUIREMENT_REFERENCE_ROOT", "./data/requirement_reference"))
TARGET_DIRS = ["강제 규정", "기술", "요구사항 가이드"]
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

DOC_TYPE_BY_DIR = {
    "강제 규정": "compliance_rule",
    "기술": "technical_guide",
    "요구사항 가이드": "requirement_writing_guide",
}


def normalize_text(value: Any) -> str:
    text = str(value if value is not None else "").replace("\n", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def split_text(text: str, chunk_size: int = 900, overlap: int = 120) -> List[str]:
    chunks = []
    start = 0
    while start < len(text):
        chunk = text[start : start + chunk_size].strip()
        if len(chunk) >= 80:
            chunks.append(chunk)
        start += chunk_size - overlap
    return chunks


def read_pdf(path: Path) -> Iterable[tuple[int, str]]:
    doc = fitz.open(path)
    try:
        for idx, page in enumerate(doc, start=1):
            text = normalize_text(page.get_text())
            if text:
                yield idx, text
    finally:
        doc.close()


def read_docx(path: Path) -> Iterable[tuple[int, str]]:
    document = docx.Document(path)
    lines = []
    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                lines.append(" | ".join(cells))
    text = "\n".join(lines)
    if text:
        yield 1, text


def read_text_file(path: Path) -> Iterable[tuple[int, str]]:
    for encoding in ["utf-8", "utf-8-sig", "cp949", "euc-kr"]:
        try:
            text = normalize_text(path.read_text(encoding=encoding))
            if text:
                yield 1, text
            return
        except UnicodeDecodeError:
            continue
    raise UnicodeError(f"지원하지 않는 인코딩입니다: {path}")


def read_document(path: Path) -> Iterable[tuple[int, str]]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return read_pdf(path)
    if ext == ".docx":
        return read_docx(path)
    if ext in {".txt", ".md"}:
        return read_text_file(path)
    raise ValueError(f"지원하지 않는 문서 형식입니다: {path}")


def iter_source_files() -> Iterable[tuple[str, Path]]:
    for dir_name in TARGET_DIRS:
        target_dir = DATA_ROOT / dir_name
        if not target_dir.exists():
            print(f"[스킵] 폴더 없음: {target_dir}")
            continue

        for path in sorted(target_dir.iterdir()):
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                yield dir_name, path


def build_payload(
    *,
    text: str,
    chunk_id: str,
    dir_name: str,
    source_file: Path,
    page: int,
    chunk_index: int,
) -> Dict[str, Any]:
    doc_type = DOC_TYPE_BY_DIR.get(dir_name, "requirement_reference")
    return {
        "text": text,
        "chunk_id": chunk_id,
        "doc_type": doc_type,
        "domain": "requirements",
        "source_name": source_file.stem,
        "section": dir_name,
        "title": f"{source_file.stem} p.{page}-{chunk_index}",
        "applies_to": "requirements_definition,requirements_validation,proposal_response",
        "priority": "required" if doc_type == "compliance_rule" else "reference",
        "source_file": source_file.name,
        "version": "",
        "chunk_type": "reference",
        "keywords": [dir_name, doc_type, "요구사항"],
        "is_active": True,
        "language": "ko",
        "page": page,
    }


def extract_payloads() -> List[Dict[str, Any]]:
    payloads = []
    for dir_name, source_file in iter_source_files():
        print(f"[처리] {dir_name} / {source_file.name}")
        for page, page_text in read_document(source_file):
            for chunk_index, chunk in enumerate(split_text(page_text), start=1):
                base = f"{dir_name}:{source_file.name}:{page}:{chunk_index}"
                chunk_id = f"requirement_reference_{uuid.uuid5(uuid.NAMESPACE_DNS, base)}"
                payloads.append(
                    build_payload(
                        text=chunk,
                        chunk_id=chunk_id,
                        dir_name=dir_name,
                        source_file=source_file,
                        page=page,
                        chunk_index=chunk_index,
                    )
                )
    return payloads


def upsert_payloads(payloads: List[Dict[str, Any]], batch_size: int = 32):
    client = get_client()
    embedder = get_embedder()

    for i in tqdm(range(0, len(payloads), batch_size)):
        batch = payloads[i : i + batch_size]
        vectors = embedder.encode(
            [p["text"] for p in batch],
            normalize_embeddings=True,
            show_progress_bar=False,
        ).tolist()
        points = [
            PointStruct(
                id=str(uuid.uuid5(uuid.NAMESPACE_DNS, payload["chunk_id"])),
                vector=vector,
                payload=payload,
            )
            for vector, payload in zip(vectors, batch)
        ]
        client.upsert(collection_name=COLLECTION_NAME, points=points)

    print(f"[적재 완료] collection={COLLECTION_NAME}, chunks={len(payloads)}")


def main():
    ensure_named_collection(COLLECTION_NAME, recreate=False)
    payloads = extract_payloads()
    print(f"[추출 완료] requirement reference chunks={len(payloads)}")
    upsert_payloads(payloads)


if __name__ == "__main__":
    main()
